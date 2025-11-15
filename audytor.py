import asyncio
import json
import re
import time
import urllib.parse
import urllib.robotparser as rps
from collections import deque, defaultdict, Counter
from typing import Dict, Any, List, Optional, Set, Tuple
from datetime import datetime
import os
import aiohttp
from bs4 import BeautifulSoup
from url_normalize import url_normalize
import tldextract

import extruct
from w3lib.html import get_base_url

try:
    from tqdm import tqdm

    HAS_TQDM = True
except:
    HAS_TQDM = False

try:
    import trafilatura

    HAS_TRAFILATURA = True
except:
    HAS_TRAFILATURA = False

try:
    from docx import Document as WordDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    HAS_DOCX = True
except Exception:
    HAS_DOCX = False
    print("⚠️  Brak biblioteki python-docx. Raport Word nie będzie wygenerowany.")

try:
    import pandas as pd
except Exception:
    pd = None

START_URL = "https://ekantor.pl/"
MAX_PAGES = 300
MAX_DEPTH = 3
TIMEOUT = 25
CONCURRENCY = 10
USER_AGENT = "SiteAuditorBot/1.0 (+https://twojadomena.example/audyt)"
RESPECT_ROBOTS = True
DOMAIN_SCOPE = "root"
USE_PAGESPEED = False
PAGESPEED_API_KEY = os.getenv("PAGESPEED_API_KEY",
                              "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")
USE_AI_SUMMARY = True
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = "gpt-4o-mini"
OUTPUT_DIR = None

EXCLUDED_PATTERNS = [
    r'/polityka[_-]prywatnosci',
    r'/privacy[_-]policy',
    r'/regulamin',
    r'/terms',
    r'/sitemap',
    r'/robots\.txt',
    r'/cookie[s]?[_-]policy',
    r'/disclaimer',
    r'/terms-of-service',
    r'/legal',
    r'^/cdn-cgi/',
    r'/cdn-cgi/l/email-protection',
]

SHOW_REMEDIATIONS = False


def same_site(u1: str, u2: str) -> bool:
    a = urllib.parse.urlparse(u1)
    b = urllib.parse.urlparse(u2)
    if DOMAIN_SCOPE == "sub":
        ea = tldextract.extract(a.netloc)
        eb = tldextract.extract(b.netloc)
        return (a.scheme in ("http", "https")) and (ea.registered_domain == eb.registered_domain)
    else:
        return (a.scheme, a.netloc) == (b.scheme, b.netloc)


def absolutize(base: str, link: str) -> str:
    return url_normalize(urllib.parse.urljoin(base, link))


def clean_text(soup: BeautifulSoup) -> str:
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text(" ", strip=True)
    return re.sub(r"\s+", " ", text)


def is_excluded_url(url: str) -> bool:
    parsed = urllib.parse.urlparse(url)
    path = parsed.path.lower()
    for pattern in EXCLUDED_PATTERNS:
        if re.search(pattern, path, re.IGNORECASE):
            return True
    return False


def calculate_meta_score(title: str, description: str) -> Dict[str, Any]:
    title_len = len(title)
    desc_len = len(description)
    title_optimal = 50 <= title_len <= 60
    title_acceptable = 30 <= title_len <= 65
    desc_optimal = 150 <= desc_len <= 160
    desc_acceptable = 120 <= desc_len <= 165
    return {
        "title_length": title_len,
        "title_optimal": title_optimal,
        "title_acceptable": title_acceptable,
        "title_too_short": title_len < 30,
        "title_too_long": title_len > 65,
        "desc_length": desc_len,
        "desc_optimal": desc_optimal,
        "desc_acceptable": desc_acceptable,
        "desc_too_short": desc_len < 120,
        "desc_too_long": desc_len > 165,
        "has_title": bool(title),
        "has_description": bool(description),
    }


def extract_nap_signals(soup: BeautifulSoup, text: str) -> Dict[str, Any]:
    phone_patterns = [
        r'\+?48\s?[\d\s\-]{9,}',
        r'\(\d{3}\)\s?\d{3}[\s\-]?\d{4}',
        r'\d{3}[\s\-]?\d{3}[\s\-]?\d{4}',
    ]
    phones = []
    for pattern in phone_patterns:
        phones.extend(re.findall(pattern, text))
    address_indicators = ['ul.', 'ulica', 'al.', 'aleja', 'street', 'avenue', 'road']
    has_address_indicators = any(ind in text.lower() for ind in address_indicators)
    schema_scripts = soup.find_all('script', type='application/ld+json')
    has_local_schema = False
    for script in schema_scripts:
        try:
            data = json.loads(script.string)
            if isinstance(data, dict):
                schema_type = data.get('@type', '')
                if isinstance(schema_type, str):
                    if any(t in schema_type for t in ['LocalBusiness', 'Organization', 'Store', 'Restaurant']):
                        has_local_schema = True
                        break
        except:
            pass
    return {
        "phone_numbers_found": len(phones),
        "has_address_indicators": has_address_indicators,
        "has_local_business_schema": has_local_schema,
        "nap_score": sum([
            len(phones) > 0,
            has_address_indicators,
            has_local_schema
        ]),
    }


def analyze_eeat_signals(soup: BeautifulSoup, text: str, url: str) -> Dict[str, Any]:
    author_indicators = ['author', 'autor', 'written by', 'by', 'redaktor']
    has_author = False
    for ind in author_indicators:
        if soup.find(attrs={'class': re.compile(ind, re.I)}) or \
                soup.find(attrs={'id': re.compile(ind, re.I)}) or \
                soup.find(attrs={'itemprop': ind}):
            has_author = True
            break
    date_indicators = ['published', 'pubdate', 'datePublished', 'article:published_time']
    has_date = False
    for ind in date_indicators:
        if soup.find(attrs={'itemprop': ind}) or soup.find('time') or soup.find('meta', property=ind):
            has_date = True
            break
    expertise_keywords = ['certyfikat', 'certificate', 'licencja', 'license', 'dyplom', 'diploma',
                          'doświadczenie', 'experience', 'lat doświadczenia', 'years of experience']
    has_expertise_signals = any(keyword in text.lower() for keyword in expertise_keywords)
    external_links = soup.find_all('a', href=True)
    external_quality_domains = ['.gov', '.edu', '.org', 'wikipedia.org']
    has_quality_sources = any(
        any(domain in link.get('href', '') for domain in external_quality_domains)
        for link in external_links
    )
    contact_indicators = ['kontakt', 'contact', 'email', 'telefon', 'phone', 'adres', 'address']
    has_contact_info = any(ind in text.lower() for ind in contact_indicators)
    has_ssl = url.startswith('https://')
    review_indicators = ['recenzja', 'review', 'opinia', 'opinion', 'rating', 'ocena']
    has_reviews = any(ind in text.lower() for ind in review_indicators)
    eeat_score = sum([
        has_author * 1.5,
        has_date * 1.0,
        has_expertise_signals * 1.5,
        has_quality_sources * 2.0,
        has_contact_info * 1.0,
        has_ssl * 1.0,
        has_reviews * 1.0,
    ])
    return {
        "has_author": has_author,
        "has_date": has_date,
        "has_expertise_signals": has_expertise_signals,
        "has_quality_external_links": has_quality_sources,
        "has_contact_info": has_contact_info,
        "has_ssl": has_ssl,
        "has_reviews": has_reviews,
        "eeat_score": round(eeat_score, 1),
        "eeat_max_score": 9.0,
        "eeat_percentage": round((eeat_score / 9.0) * 100, 1),
    }


def analyze_security_headers(headers: Dict[str, str], url: str, html: str = "") -> Dict[str, Any]:
    headers_lower = {k.lower(): v for k, v in headers.items()}
    security_checks = {
        "hsts": {"name": "HTTP Strict Transport Security (HSTS)", "header": "strict-transport-security",
                 "present": False, "value": "", "score": 0, "max_score": 15, "severity": "high", "description": ""},
        "x_frame_options": {"name": "X-Frame-Options", "header": "x-frame-options", "present": False, "value": "",
                            "score": 0, "max_score": 10, "severity": "high", "description": ""},
        "x_content_type_options": {"name": "X-Content-Type-Options", "header": "x-content-type-options",
                                   "present": False, "value": "", "score": 0, "max_score": 10, "severity": "medium",
                                   "description": ""},
        "content_security_policy": {"name": "Content-Security-Policy (CSP)", "header": "content-security-policy",
                                    "present": False, "value": "", "score": 0, "max_score": 20, "severity": "high",
                                    "description": ""},
        "x_xss_protection": {"name": "X-XSS-Protection", "header": "x-xss-protection", "present": False, "value": "",
                             "score": 0, "max_score": 5, "severity": "low", "description": ""},
        "referrer_policy": {"name": "Referrer-Policy", "header": "referrer-policy", "present": False, "value": "",
                            "score": 0, "max_score": 10, "severity": "medium", "description": ""},
        "permissions_policy": {"name": "Permissions-Policy", "header": "permissions-policy", "present": False,
                               "value": "", "score": 0, "max_score": 10, "severity": "low", "description": ""},
    }
    total_score = 0
    max_possible_score = sum(check["max_score"] for check in security_checks.values())
    for key, check in security_checks.items():
        header_name = check["header"]
        if header_name in headers_lower:
            check["present"] = True
            check["value"] = headers_lower[header_name]
            check["score"] = check["max_score"]
        total_score += check["score"]
    has_ssl = url.startswith('https://')
    ssl_score = 20 if has_ssl else 0
    max_possible_score += 20
    total_score += ssl_score
    has_mixed_content = False
    if has_ssl and html:
        http_resources = re.findall(r'src=["\']http://[^"\']+["\']', html, re.I)
        http_resources += re.findall(r'href=["\']http://[^"\']+\.(?:css|js)["\']', html, re.I)
        has_mixed_content = len(http_resources) > 0
    server_header = headers_lower.get('server', '')
    exposes_server_info = bool(server_header and server_header.lower() not in ['', 'cloudflare'])
    powered_by = headers_lower.get('x-powered-by', '')
    exposes_tech_stack = bool(powered_by)
    security_percentage = round((total_score / max_possible_score) * 100, 1)
    if security_percentage >= 90:
        security_level = "Excellent";
        security_emoji = "🟢"
    elif security_percentage >= 70:
        security_level = "Good";
        security_emoji = "🟡"
    elif security_percentage >= 50:
        security_level = "Fair";
        security_emoji = "🟠"
    else:
        security_level = "Poor";
        security_emoji = "🔴"
    missing_critical = [
        check["name"] for check in security_checks.values()
        if not check["present"] and check["severity"] in ["high", "medium"]
    ]
    return {
        "security_checks": security_checks,
        "has_ssl": has_ssl,
        "ssl_score": ssl_score,
        "has_mixed_content": has_mixed_content,
        "exposes_server_info": exposes_server_info,
        "server_header": server_header,
        "exposes_tech_stack": exposes_tech_stack,
        "powered_by_header": powered_by,
        "total_score": total_score,
        "max_score": max_possible_score,
        "security_percentage": security_percentage,
        "security_level": security_level,
        "security_emoji": security_emoji,
        "missing_critical": missing_critical,
        "headers_count": len([c for c in security_checks.values() if c["present"]]),
        "total_headers": len(security_checks),
    }


def generate_ai_summary(summary: Dict[str, Any], issues: Dict[str, Any]) -> str:
    if not USE_AI_SUMMARY or not OPENAI_API_KEY:
        return ""
    try:
        try:
            from openai import OpenAI
        except ImportError:
            print("⚠️  Brak biblioteki openai. Zainstaluj: pip install openai")
            return ""
        client = OpenAI(api_key=OPENAI_API_KEY)
        audit_data = {
            "url": summary["start_url"],
            "pages_analyzed": summary["pages_analyzed"],
            "seo": {
                "errors": summary["pages_with_errors"],
                "missing_title": summary["missing_title"],
                "missing_description": summary["missing_description"],
                "duplicate_titles": summary["duplicate_titles"],
                "duplicate_descriptions": summary["duplicate_descriptions"],
                "missing_canonical": summary["missing_canonical"],
                "mobile_percentage": summary["mobile_percentage"],
                "avg_eeat_score": summary["avg_eeat_score"],
            },
            "security": {
                "avg_score": summary["avg_security_score"],
                "ssl_percentage": summary["ssl_percentage"],
                "pages_poor_security": summary["pages_poor_security"],
            },
            "geo": {
                "pages_with_schema": summary["pages_with_schema"],
                "pages_without_schema": summary["pages_without_schema"],
            }
        }
        prompt = f"""Jesteś ekspertem SEO/GEO/Security. Przeanalizuj wyniki audytu i wygeneruj krótkie podsumowanie.

DANE AUDYTU:
{json.dumps(audit_data, indent=2, ensure_ascii=False)}

Wygeneruj podsumowanie (MAX 300 słów) zawierające:
1. OGÓLNA OCENA (1-2 zdania)
2. TOP 3 PRIORYTETY (lista) – konkretne action items, ale BEZ instrukcji wdrożeniowych i konfiguracji serwera
3. MOCNE STRONY (2-3 punkty)
4. OSTRZEŻENIA (jeśli są krytyczne problemy)
5. REKOMENDACJA BIZNESOWA (1 zdanie)

Styl: zwięzły, profesjonalny, emoji (🔴🟠🟡✅).
Język: polski
Format: Markdown, nagłówki ###
Nie podawaj konfiguracji serwerów (Apache/Nginx/Cloudflare) ani komend."""
        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system",
                 "content": "Jesteś ekspertem SEO, GEO i cyberbezpieczeństwa. Tworzysz zwięzłe, actionable podsumowania audytów bez instrukcji wdrożeniowych."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=800
        )
        ai_summary = response.choices[0].message.content.strip()
        return ai_summary
    except Exception as e:
        print(f"⚠️  Błąd generowania AI Summary: {e}")
        return ""


async def fetch(session: aiohttp.ClientSession, url: str) -> Tuple[Optional[int], str, str, str, Dict[str, str]]:
    try:
        async with session.get(url, allow_redirects=True, timeout=TIMEOUT, headers={"User-Agent": USER_AGENT}) as r:
            ct = r.headers.get("Content-Type", "")
            txt = await r.text(errors="ignore")
            return r.status, str(r.url), ct, txt, {k: v for k, v in r.headers.items()}
    except Exception as e:
        return None, url, "", f"__ERROR__:{e}", {}


def parse_page(html: str, url: str) -> Dict[str, Any]:
    soup = BeautifulSoup(html, "lxml")
    title = (soup.title.string.strip() if soup.title and soup.title.string else "")
    desc = ""
    mr = soup.find("meta", attrs={"name": "description"})
    if mr and mr.get("content"):
        desc = mr["content"].strip()
    viewport = soup.find("meta", attrs={"name": "viewport"})
    has_viewport = bool(viewport)
    viewport_content = viewport.get("content", "") if viewport else ""
    robots_meta = ""
    mrobots = soup.find("meta", attrs={"name": "robots"})
    if mrobots and mrobots.get("content"):
        robots_meta = mrobots["content"].lower()
    canonical = ""
    link_canon = soup.find("link", rel=lambda v: v and "canonical" in v)
    if link_canon and link_canon.get("href"):
        canonical = absolutize(url, link_canon["href"])
    h1 = [h.get_text(strip=True) for h in soup.find_all("h1")]
    h2 = [h.get_text(strip=True) for h in soup.find_all("h2")]
    h3 = [h.get_text(strip=True) for h in soup.find_all("h3")]
    imgs = soup.find_all("img")
    img_without_alt = sum(1 for i in imgs if not i.get("alt"))
    img_total = len(imgs)
    a_tags = soup.find_all("a", href=True)
    links = [absolutize(url, a["href"]) for a in a_tags if not a["href"].startswith("javascript:")]
    og_data = {}
    og_tags = soup.find_all("meta", property=re.compile(r"^og:"))
    for tag in og_tags:
        prop = tag.get("property", "");
        content = tag.get("content", "")
        if prop and content:
            og_data[prop] = content
    twitter_data = {}
    twitter_tags = soup.find_all("meta", attrs={"name": re.compile(r"^twitter:")})
    for tag in twitter_tags:
        name = tag.get("name", "");
        content = tag.get("content", "")
        if name and content:
            twitter_data[name] = content
    try:
        structured = extruct.extract(
            html,
            base_url=get_base_url(html, url),
            syntaxes=["json-ld", "microdata", "opengraph"],
            uniform=True
        )
    except Exception:
        structured = {"json-ld": [], "microdata": [], "opengraph": []}
    jsonld_types = []
    for node in structured.get("json-ld", []):
        t = node.get("@type")
        if isinstance(t, list):
            jsonld_types += t
        elif t:
            jsonld_types.append(t)
    jsonld_types = list(dict.fromkeys(jsonld_types))
    text = clean_text(soup)
    text_len = len(text)
    nap_signals = extract_nap_signals(soup, text)
    eeat_signals = analyze_eeat_signals(soup, text, url)
    meta_scores = calculate_meta_score(title, desc)
    geo_signals = {
        "has_faq_schema": "FAQPage" in jsonld_types,
        "has_article_schema": any(t in jsonld_types for t in ("Article", "NewsArticle", "BlogPosting")),
        "has_org_schema": any(t in jsonld_types for t in ("Organization", "LocalBusiness")),
        "has_breadcrumbs": "BreadcrumbList" in jsonld_types,
        "has_review_schema": any(t in jsonld_types for t in ("Review", "AggregateRating")),
        "has_product_schema": "Product" in jsonld_types,
        "clear_hierarchy": len(h1) == 1 and len(h2) > 0,
        "sufficient_text": text_len >= 1200,
        "has_navigation_schema": any('navigation' in str(s).lower() for s in soup.find_all(['nav'])),
    }
    return {
        "title": title,
        "meta_description": desc,
        "robots_meta": robots_meta,
        "canonical": canonical,
        "h1": h1,
        "h2": h2,
        "h3": h3,
        "h1_count": len(h1),
        "h2_count": len(h2),
        "h3_count": len(h3),
        "img_total": img_total,
        "img_without_alt": img_without_alt,
        "img_alt_ratio": round((img_total - img_without_alt) / max(1, img_total) * 100, 1),
        "has_viewport": has_viewport,
        "viewport_content": viewport_content,
        "is_mobile_friendly": has_viewport and 'width=device-width' in viewport_content.lower(),
        "opengraph": og_data,
        "twitter_cards": twitter_data,
        "has_og_image": "og:image" in og_data,
        "has_og_title": "og:title" in og_data,
        "has_og_description": "og:description" in og_data,
        "has_twitter_card": "twitter:card" in twitter_data,
        "jsonld_types": jsonld_types,
        "schema_count": len(jsonld_types),
        "text_len": text_len,
        "word_count": len(text.split()),
        "links": links,
        "internal_links": len([l for l in links if same_site(url, l)]),
        "external_links": len([l for l in links if not same_site(url, l)]),
        "nap_signals": nap_signals,
        "eeat_signals": eeat_signals,
        "geo_signals": geo_signals,
        "meta_scores": meta_scores,
        "is_excluded": is_excluded_url(url),
    }


async def check_pagespeed(url: str) -> Dict[str, Any]:
    if not USE_PAGESPEED or not PAGESPEED_API_KEY:
        return {}
    try:
        api_url = f"https://www.googleapis.com/pagespeedonline/v5/runPagespeed"
        params = {
            "url": url,
            "key": PAGESPEED_API_KEY,
            "strategy": "mobile",
            "category": ["performance", "accessibility", "best-practices", "seo"]
        }
        async with aiohttp.ClientSession() as session:
            async with session.get(api_url, params=params, timeout=60) as resp:
                if resp.status == 200:
                    data = await resp.json()
                    lighthouse = data.get("lighthouseResult", {})
                    categories = lighthouse.get("categories", {})
                    audits = lighthouse.get("audits", {})
                    cwv = {
                        "lcp": audits.get("largest-contentful-paint", {}).get("numericValue"),
                        "cls": audits.get("cumulative-layout-shift", {}).get("numericValue"),
                        "fid": audits.get("max-potential-fid", {}).get("numericValue"),
                    }
                    return {
                        "performance_score": categories.get("performance", {}).get("score", 0) * 100,
                        "accessibility_score": categories.get("accessibility", {}).get("score", 0) * 100,
                        "seo_score": categories.get("seo", {}).get("score", 0) * 100,
                        "core_web_vitals": cwv,
                    }
    except Exception as e:
        print(f"⚠️  PageSpeed API error for {url}: {e}")
    return {}


async def build_robots(session: aiohttp.ClientSession, root: str) -> rps.RobotFileParser:
    rp = rps.RobotFileParser()
    p = urllib.parse.urlparse(root)
    robots_url = f"{p.scheme}://{p.netloc}/robots.txt"
    status, final, ct, text, headers = await fetch(session, robots_url)
    rp.set_url(robots_url)
    if status and status < 400 and isinstance(text, str) and not text.startswith("__ERROR__"):
        rp.parse(text.splitlines())
    else:
        rp.parse([])
    return rp


async def discover_sitemaps(session: aiohttp.ClientSession, root: str) -> List[str]:
    p = urllib.parse.urlparse(root)
    candidates = [
        f"{p.scheme}://{p.netloc}/sitemap.xml",
        f"{p.scheme}://{p.netloc}/sitemap_index.xml",
    ]
    found = []
    for u in candidates:
        status, final, ct, text, headers = await fetch(session, u)
        if status and status < 400 and "xml" in ct.lower():
            found.append(final)
    return found


def parse_sitemap_xml(xml_text: str) -> List[str]:
    urls = re.findall(r"<loc>(.*?)</loc>", xml_text, flags=re.I | re.S)
    return [u.strip() for u in urls if u.strip()]


async def fetch_and_parse_sitemaps(session: aiohttp.ClientSession, sitemap_urls: List[str]) -> List[str]:
    urls: List[str] = []
    for sm in sitemap_urls:
        status, final, ct, text, headers = await fetch(session, sm)
        if status and status < 400 and isinstance(text, str) and not text.startswith("__ERROR__"):
            urls += parse_sitemap_xml(text)
    return list(dict.fromkeys(urls))


async def crawl(start_url: str) -> Dict[str, Any]:
    q = deque([(start_url, 0)])
    seen: Set[str] = {start_url}
    results: Dict[str, Any] = {}
    pbar = tqdm(total=MAX_PAGES, desc="Crawling", unit="page") if HAS_TQDM else None
    async with aiohttp.ClientSession() as session:
        rp = await build_robots(session, start_url) if RESPECT_ROBOTS else None
        sem = asyncio.Semaphore(CONCURRENCY)
        try:
            sitemaps = await discover_sitemaps(session, start_url)
            if sitemaps:
                urls_from_sm = await fetch_and_parse_sitemaps(session, sitemaps)
                for u in urls_from_sm[:max(0, MAX_PAGES - len(q))]:
                    if u not in seen and same_site(start_url, u) and not is_excluded_url(u):
                        seen.add(u)
                        q.append((u, 1))
        except Exception as e:
            print(f"⚠️  Błąd przy pobieraniu sitemap: {e}")

        async def worker():
            while q and len(results) < MAX_PAGES:
                url, depth = q.popleft()
                if RESPECT_ROBOTS and rp and not rp.can_fetch(USER_AGENT, url):
                    results[url] = {"url": url, "error": "blocked_by_robots"}
                    if pbar: pbar.update(1)
                    continue
                async with sem:
                    status, final, ct, html, headers = await fetch(session, url)
                    item: Dict[str, Any] = {"url": url, "final_url": final, "status": status, "content_type": ct,
                                            "headers": headers}
                    if not status or (isinstance(html, str) and html.startswith("__ERROR__")):
                        item["error"] = html if isinstance(html, str) else "fetch_error"
                        results[url] = item
                        if pbar: pbar.update(1)
                        continue
                    if ct and "text/html" in ct:
                        parsed = parse_page(html, final)
                        item.update(parsed)
                        security_analysis = analyze_security_headers(headers, final, html)
                        item["security"] = security_analysis
                        if USE_PAGESPEED and len(results) < 5:
                            item["pagespeed"] = await check_pagespeed(final)
                        for link in item.get("links", []):
                            if not same_site(start_url, link):
                                continue
                            if is_excluded_url(link):
                                continue
                            if link not in seen and depth + 1 <= MAX_DEPTH:
                                seen.add(link)
                                q.append((link, depth + 1))
                    else:
                        item["note"] = "Pominięto (non-HTML)"
                    results[url] = item
                    if pbar: pbar.update(1)

        tasks = [asyncio.create_task(worker()) for _ in range(CONCURRENCY)]
        await asyncio.gather(*tasks)
    if pbar: pbar.close()
    return results


def find_duplicates(all_pages: Dict[str, Any]) -> Dict[str, List]:
    title_map = defaultdict(list)
    desc_map = defaultdict(list)
    for url, data in all_pages.items():
        if data.get('is_excluded'):
            continue
        title = data.get('title', '').strip()
        desc = data.get('meta_description', '').strip()
        if title:
            title_map[title].append(url)
        if desc:
            desc_map[desc].append(url)
    duplicates = {
        "title": {k: v for k, v in title_map.items() if len(v) > 1},
        "description": {k: v for k, v in desc_map.items() if len(v) > 1},
    }
    return duplicates


def analyze_issues(all_pages: Dict[str, Any]) -> Dict[str, Any]:
    issues = {
        "critical_errors": [],
        "missing_title": [],
        "missing_description": [],
        "missing_canonical": [],
        "missing_h1": [],
        "multiple_h1": [],
        "images_no_alt": [],
        "no_viewport": [],
        "no_og_tags": [],
        "no_twitter_cards": [],
        "missing_schema": [],
        "weak_eeat": [],
        "poor_local_seo": [],
        "thin_content": [],
        "title_issues": [],
        "description_issues": [],
        "no_ssl": [],
        "missing_security_headers": [],
        "poor_security": [],
        "mixed_content": [],
        "info_disclosure": [],
    }
    for url, data in all_pages.items():
        if data.get('is_excluded'):
            continue
        ct = data.get('content_type', '') or ''
        status = data.get('status')
        if data.get('error') or (status and 400 <= status < 500):
            if 'text/html' in ct or not ct:
                issues['critical_errors'].append({'url': url, 'status': status, 'error': data.get('error', '')})
        if not status or status >= 400:
            continue
        if not data.get('title'):
            issues['missing_title'].append(url)
        else:
            meta_scores = data.get('meta_scores', {})
            if meta_scores.get('title_too_short') or meta_scores.get('title_too_long'):
                issues['title_issues'].append({
                    'url': url,
                    'title': data.get('title'),
                    'length': meta_scores.get('title_length'),
                    'too_short': meta_scores.get('title_too_short'),
                    'too_long': meta_scores.get('title_too_long'),
                })
        if not data.get('meta_description'):
            issues['missing_description'].append(url)
        else:
            meta_scores = data.get('meta_scores', {})
            if meta_scores.get('desc_too_short') or meta_scores.get('desc_too_long'):
                issues['description_issues'].append({
                    'url': url,
                    'description': data.get('meta_description')[:100],
                    'length': meta_scores.get('desc_length'),
                    'too_short': meta_scores.get('desc_too_short'),
                    'too_long': meta_scores.get('desc_too_long'),
                })
        if not data.get('canonical'):
            issues['missing_canonical'].append(url)
        h1_count = data.get('h1_count', 0)
        if h1_count == 0:
            issues['missing_h1'].append(url)
        elif h1_count > 1:
            issues['multiple_h1'].append({'url': url, 'h1_count': h1_count, 'h1_list': data.get('h1', [])})
        if data.get('img_without_alt', 0) > 0:
            issues['images_no_alt'].append({
                'url': url,
                'missing_alt': data.get('img_without_alt'),
                'total_images': data.get('img_total'),
                'alt_ratio': data.get('img_alt_ratio'),
            })
        if not data.get('is_mobile_friendly'):
            issues['no_viewport'].append(url)
        if not data.get('has_og_image') or not data.get('has_og_title'):
            issues['no_og_tags'].append({
                'url': url,
                'has_og_image': data.get('has_og_image'),
                'has_og_title': data.get('has_og_title'),
                'has_og_description': data.get('has_og_description'),
            })
        if not data.get('has_twitter_card'):
            issues['no_twitter_cards'].append(url)
        if data.get('schema_count', 0) == 0:
            issues['missing_schema'].append(url)
        eeat = data.get('eeat_signals', {})
        if eeat.get('eeat_percentage', 100) < 50:
            issues['weak_eeat'].append({
                'url': url,
                'eeat_score': eeat.get('eeat_score'),
                'eeat_percentage': eeat.get('eeat_percentage'),
                'missing': [k for k, v in eeat.items() if k.startswith('has_') and not v]
            })
        nap = data.get('nap_signals', {})
        if nap.get('nap_score', 0) < 2:
            issues['poor_local_seo'].append({
                'url': url,
                'nap_score': nap.get('nap_score'),
                'phone_numbers': nap.get('phone_numbers_found'),
                'has_address': nap.get('has_address_indicators'),
                'has_local_schema': nap.get('has_local_business_schema'),
            })
        word_count = data.get('word_count', 0) or 0
        if word_count < 300 and word_count > 0:
            issues['thin_content'].append({'url': url, 'word_count': word_count, 'text_len': data.get('text_len', 0)})
        security = data.get('security', {})
        if not security.get('has_ssl'):
            issues['no_ssl'].append(url)
        sec_percentage = security.get('security_percentage', 100)
        if sec_percentage < 50:
            issues['poor_security'].append({
                'url': url,
                'security_percentage': sec_percentage,
                'security_level': security.get('security_level'),
                'missing_headers': security.get('missing_critical', []),
            })
        headers_count = security.get('headers_count', 0)
        if headers_count < 3:
            issues['missing_security_headers'].append({
                'url': url,
                'headers_count': headers_count,
                'missing_critical': security.get('missing_critical', []),
            })
        if security.get('has_mixed_content'):
            issues['mixed_content'].append(url)
        if security.get('exposes_server_info') or security.get('exposes_tech_stack'):
            issues['info_disclosure'].append({
                'url': url,
                'server_header': security.get('server_header'),
                'powered_by': security.get('powered_by_header'),
            })
    return issues


def calculate_overall_score(summary: Dict[str, Any]) -> Tuple[int, str]:
    pages = max(1, summary["pages_analyzed"])
    availability = summary["pages_ok"] / pages
    meta_ok_pages = pages - (summary["missing_title"] + summary["missing_description"])
    meta_quality = max(0.0, (meta_ok_pages - 0.25 * (summary["title_issues"] + summary["description_issues"])) / pages)
    mobile = summary["mobile_percentage"] / 100.0
    schema = (summary["pages_with_schema"] / pages) if pages else 0.0
    eeat = summary["avg_eeat_score"] / 100.0
    security = summary["avg_security_score"] / 100.0
    if summary["pages_no_ssl"] > 0:
        security = max(0.0, security - 0.10)
    W = {"availability": 0.30, "meta": 0.15, "mobile": 0.15, "schema": 0.10, "eeat": 0.10, "security": 0.20}
    score = (availability * W["availability"] + meta_quality * W["meta"] + mobile * W["mobile"] + schema * W[
        "schema"] + eeat * W["eeat"] + security * W["security"]) * 100.0
    score_int = max(0, min(100, int(round(score))))
    if score_int >= 85:
        grade = "Excellent"
    elif score_int >= 70:
        grade = "Good"
    elif score_int >= 50:
        grade = "Fair"
    else:
        grade = "Poor"
    return score_int, grade


def calculate_summary(all_pages: Dict[str, Any], issues: Dict[str, Any], duplicates: Dict) -> Dict[str, Any]:
    analyzed_pages = {url: data for url, data in all_pages.items() if not data.get('is_excluded')}
    excluded_count = len(all_pages) - len(analyzed_pages)
    pages_with_errors = len(issues['critical_errors'])
    pages_ok = len([p for p in analyzed_pages.values() if p.get('status') == 200])
    mobile_friendly = sum(1 for p in analyzed_pages.values() if p.get('is_mobile_friendly'))
    mobile_percentage = round(mobile_friendly / max(1, len(analyzed_pages)) * 100, 1)
    pages_with_schema = sum(1 for p in analyzed_pages.values() if p.get('schema_count', 0) > 0)
    avg_schema_types = sum(p.get('schema_count', 0) for p in analyzed_pages.values()) / max(1, len(analyzed_pages))
    avg_eeat = sum(p.get('eeat_signals', {}).get('eeat_percentage', 0) for p in analyzed_pages.values()) / max(1,
                                                                                                               len(analyzed_pages))
    local_optimized = sum(1 for p in analyzed_pages.values() if p.get('nap_signals', {}).get('nap_score', 0) >= 2)
    avg_security = sum(p.get('security', {}).get('security_percentage', 0) for p in analyzed_pages.values()) / max(1,
                                                                                                                   len(analyzed_pages))
    pages_with_ssl = sum(1 for p in analyzed_pages.values() if p.get('security', {}).get('has_ssl'))
    ssl_percentage = round(pages_with_ssl / max(1, len(analyzed_pages)) * 100, 1)
    result = {
        "start_url": START_URL,
        "pages_crawled": len(all_pages),
        "pages_analyzed": len(analyzed_pages),
        "pages_excluded": excluded_count,
        "pages_ok": pages_ok,
        "pages_with_errors": pages_with_errors,
        "missing_title": len(issues['missing_title']),
        "missing_description": len(issues['missing_description']),
        "title_issues": len(issues['title_issues']),
        "description_issues": len(issues['description_issues']),
        "duplicate_titles": len(duplicates['title']),
        "duplicate_descriptions": len(duplicates['description']),
        "missing_canonical": len(issues['missing_canonical']),
        "missing_h1": len(issues['missing_h1']),
        "multiple_h1": len(issues['multiple_h1']),
        "pages_with_alt_issues": len(issues['images_no_alt']),
        "total_images_without_alt": sum(i['missing_alt'] for i in issues['images_no_alt']),
        "mobile_friendly_pages": mobile_friendly,
        "mobile_percentage": mobile_percentage,
        "pages_without_viewport": len(issues['no_viewport']),
        "pages_without_og": len(issues['no_og_tags']),
        "pages_without_twitter": len(issues['no_twitter_cards']),
        "pages_with_schema": pages_with_schema,
        "pages_without_schema": len(issues['missing_schema']),
        "avg_schema_types": round(avg_schema_types, 1),
        "avg_eeat_score": round(avg_eeat, 1),
        "pages_weak_eeat": len(issues['weak_eeat']),
        "local_optimized_pages": local_optimized,
        "pages_poor_local_seo": len(issues['poor_local_seo']),
        "thin_content_pages": len(issues['thin_content']),
        "avg_security_score": round(avg_security, 1),
        "pages_with_ssl": pages_with_ssl,
        "ssl_percentage": ssl_percentage,
        "pages_no_ssl": len(issues['no_ssl']),
        "pages_poor_security": len(issues['poor_security']),
        "pages_missing_security_headers": len(issues['missing_security_headers']),
        "pages_with_mixed_content": len(issues['mixed_content']),
        "pages_with_info_disclosure": len(issues['info_disclosure']),
        "generated_at": datetime.now().strftime("%Y-%m-%d"),
    }
    overall_score, overall_grade = calculate_overall_score(result)
    result["overall_score"] = overall_score
    result["overall_grade"] = overall_grade
    return result


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color');
    color.set(qn('w:val'), '0563C1');
    rPr.append(color)
    u = OxmlElement('w:u');
    u.set(qn('w:val'), 'single');
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def create_word_report(all_pages: Dict[str, Any], summary: Dict[str, Any],
                       issues: Dict[str, Any], duplicates: Dict, word_path: str):
    if not HAS_DOCX:
        print("⚠️  Pomijam generowanie raportu Word")
        return
    doc = WordDocument()
    title = doc.add_heading('Audyt SEO/AEO/GEO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(31, 71, 136)
    title_run.font.size = Pt(36)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{summary['start_url']}\n")
    run.font.size = Pt(18);
    run.font.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Data audytu: {summary['generated_at']}")
    run.font.size = Pt(11);
    run.italic = True;
    run.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_paragraph()

    # WYNIK AUDYTU z tabelką wyjaśniającą scoring
    box = doc.add_paragraph()
    run = box.add_run(f"WYNIK AUDYTU: {summary['overall_score']}/100  ({summary['overall_grade']})")
    run.font.size = Pt(26);
    run.bold = True;
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph()

    # Dodanie tabelki wyjaśniającej scoring
    doc.add_heading('Skład oceny ogólnej:', level=3)
    scoring_table = doc.add_table(rows=1, cols=3)
    scoring_table.style = 'Light Grid Accent 1'
    hdr = scoring_table.rows[0].cells
    hdr[0].text = 'Kategoria'
    hdr[1].text = 'Waga'
    hdr[2].text = 'Opis'

    scoring_components = [
        ('Dostępność (HTTP 200)', '30%', 'Procent stron bez błędów HTTP (4xx'),
        ('Meta tagi', '15%', 'Obecność i jakość Title oraz Description'),
        ('Mobile-friendly', '15%', 'Responsywność i meta viewport'),
        ('Schema.org', '10%', 'Dane strukturalne JSON-LD'),
        ('E-E-A-T', '10%', 'Eksperckość, Autorytet, Zaufanie'),
        ('Bezpieczeństwo', '20%', 'HTTPS, Security Headers, brak mixed content'),
    ]

    for category, weight, description in scoring_components:
        row = scoring_table.add_row().cells
        row[0].text = category
        row[1].text = weight
        row[2].text = description

    doc.add_paragraph()

    # Najważniejsze problemy z licznikami X/Y i opisami
    doc.add_heading('⚠️ Najważniejsze problemy', level=2)

    # Słownik z opisami problemów
    problem_descriptions = {
        "pages_with_errors": "Strony z kodami błędów HTTP (4xx, 5xx). Uniemożliwiają indeksację i powodują utratę ruchu.",
        "pages_no_ssl": "Brak certyfikatu SSL/HTTPS. Obniża zaufanie użytkowników i ranking w Google.",
        "missing_title": "Każda strona potrzebuje unikalnego tagu <title> (50-60 znaków) dla lepszej widoczności w wynikach wyszukiwania.",
        "missing_description": "Meta description (150-160 znaków) to pierwszy kontakt użytkownika z Twoją stroną w wynikach Google.",
        "missing_canonical": "Tag canonical zapobiega problemom z duplikacją treści i pomaga Google wybrać właściwą wersję strony.",
        "pages_without_viewport": "Meta viewport to podstawa responsywności. Bez niego strona źle wyświetla się na urządzeniach mobilnych.",
        "pages_without_schema": "Dane strukturalne (Schema.org) pomagają Google lepiej zrozumieć zawartość i wyświetlać rich snippets.",
        "pages_poor_security": "Słabe zabezpieczenia (<50%). Brak security headers naraża użytkowników i obniża zaufanie do witryny.",
        "thin_content_pages": "Strony z mniej niż 300 słowami. Google preferuje wartościowe, szczegółowe treści.",
        "pages_weak_eeat": "Słabe sygnały E-E-A-T (<50%). Dodaj autora, datę publikacji, certyfikaty i linki do wiarygodnych źródeł.",
    }

    top_issues = [
        ("Błędy HTTP (4xx/5xx)", summary["pages_with_errors"], summary["pages_analyzed"], "pages_with_errors"),
        ("Brak SSL (HTTPS)", summary["pages_no_ssl"], summary["pages_analyzed"], "pages_no_ssl"),
        ("Brak Title", summary["missing_title"], summary["pages_analyzed"], "missing_title"),
        ("Brak Meta Description", summary["missing_description"], summary["pages_analyzed"], "missing_description"),
        ("Brak canonical", summary["missing_canonical"], summary["pages_analyzed"], "missing_canonical"),
        ("Brak meta viewport (mobile)", summary["pages_without_viewport"], summary["pages_analyzed"],
         "pages_without_viewport"),
        ("Brak Schema.org", summary["pages_without_schema"], summary["pages_analyzed"], "pages_without_schema"),
        ("Słabe bezpieczeństwo (<50%)", summary["pages_poor_security"], summary["pages_analyzed"],
         "pages_poor_security"),
        ("Cienka treść (<300 słów)", summary["thin_content_pages"], summary["pages_analyzed"], "thin_content_pages"),
        ("Słabe E-E-A-T (<50%)", summary["pages_weak_eeat"], summary["pages_analyzed"], "pages_weak_eeat"),
    ]

    for label, problem_count, total_pages, key in sorted(top_issues, key=lambda kv: kv[1], reverse=True)[:6]:
        if problem_count > 0:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"{label}: {problem_count}/{total_pages}\n")
            run.bold = True
            if key in problem_descriptions:
                run2 = p.add_run(problem_descriptions[key])
                run2.font.size = Pt(10)
                run2.italic = True
                run2.font.color.rgb = RGBColor(80, 80, 80)

    # USUNIĘTO notatkę o wykluczonych adresach

    doc.add_page_break()

    doc.add_heading('📋 Spis Treści', 1)
    toc_items = [
        "1. Executive Summary - Kluczowe Liczby",
        "2. Priorytety (od krytycznych do lekkich)",
        "3. Analiza Meta Tagów",
        "4. Analiza Techniczna SEO",
        "5. Mobilność i Responsywność",
        "6. Open Graph i Twitter Cards",
        "7. Dane Strukturalne (Schema.org)",
        "8. E-E-A-T",
        "9. Local SEO (NAP)",
        "10. Jakość Treści",
        "11. BEZPIECZEŃSTWO (Security Headers)",
        "12. Legenda i objaśnienia",
        "13. AI-Powered Executive Summary",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
    doc.add_page_break()

    doc.add_heading('1. 📊 Executive Summary - Kluczowe Liczby', 1)
    doc.add_heading('📌 Szybkie Podsumowanie SEO', 2)
    p = doc.add_paragraph();
    run = p.add_run('Stan strony w pigułce:\n')
    run.font.size = Pt(11);
    run.bold = True

    # Problemy KRYTYCZNE: 4xx, brak SSL, brak Title, brak Description
    seo_critical = summary['pages_with_errors'] + summary['pages_no_ssl'] + summary['missing_title'] + summary[
        'missing_description']

    # Ostrzeżenia: problemy z długością, canonical, schema, itp.
    seo_warnings = summary['title_issues'] + summary['description_issues'] + summary['missing_canonical'] + summary[
        'pages_without_schema']

    if seo_critical == 0 and seo_warnings < 10:
        seo_status = "✅ Świetnie!";
        seo_color = RGBColor(0, 150, 0);
        seo_text = "Drobne optymalizacje."
    elif seo_critical < 5 and seo_warnings < 30:
        seo_status = "⚠️ Do poprawy";
        seo_color = RGBColor(200, 100, 0);
        seo_text = f"Problemy krytyczne: {seo_critical} | Ostrzeżenia: {seo_warnings}"
    else:
        seo_status = "🔴 Wymaga uwagi!";
        seo_color = RGBColor(200, 0, 0);
        seo_text = f"Problemy krytyczne: {seo_critical} (błędy 4xx, brak SSL/Title/Description) | Ostrzeżenia: {seo_warnings} (długość meta, canonical, schema)"

    p = doc.add_paragraph()
    run = p.add_run(f"Status SEO: {seo_status}\n")
    run.font.size = Pt(14);
    run.font.bold = True;
    run.font.color.rgb = seo_color
    run = p.add_run(seo_text);
    run.font.size = Pt(11)
    doc.add_paragraph()
    p = doc.add_paragraph();
    run = p.add_run('🎯 Top 3 Priorytety:\n');
    run.font.size = Pt(11);
    run.bold = True
    priorities = []

    total_pages = summary['pages_analyzed']

    if summary['pages_with_errors'] > 0:
        pct = round((summary['pages_with_errors'] / total_pages) * 100, 1)
        priorities.append({
            'text': f"🔴 Napraw {summary['pages_with_errors']}/{total_pages} stron z błędami HTTP 4xx ({pct}%)",
            'desc': "Strony niedostępne dla użytkowników i botów Google",
            'count': summary['pages_with_errors']
        })

    if summary['pages_no_ssl'] > 0:
        pct = round((summary['pages_no_ssl'] / total_pages) * 100, 1)
        priorities.append({
            'text': f"🔴 Wymuś pełny HTTPS na {summary['pages_no_ssl']}/{total_pages} stronach ({pct}%)",
            'desc': "Brak szyfrowania obniża ranking i zaufanie",
            'count': summary['pages_no_ssl']
        })

    if summary['missing_title'] > 0:
        pct = round((summary['missing_title'] / total_pages) * 100, 1)
        priorities.append({
            'text': f"🔴 Dodaj Title do {summary['missing_title']}/{total_pages} stron ({pct}%)",
            'desc': "Brak tytułu = niewidoczność w wynikach Google",
            'count': summary['missing_title']
        })

    if summary['missing_description'] > 0 and len(priorities) < 3:
        pct = round((summary['missing_description'] / total_pages) * 100, 1)
        priorities.append({
            'text': f"🔴 Dodaj Meta Description do {summary['missing_description']}/{total_pages} stron ({pct}%)",
            'desc': "Wpływa na CTR (Click-Through Rate) z wyszukiwarki",
            'count': summary['missing_description']
        })

    if summary['pages_without_schema'] > 0 and len(priorities) < 3:
        pct = round((summary['pages_without_schema'] / total_pages) * 100, 1)
        priorities.append({
            'text': f"🟠 Dodaj Schema.org do {summary['pages_without_schema']}/{total_pages} stron ({pct}%)",
            'desc': "Brak rich snippets w Google (gwiazdki, FAQ, breadcrumbs)",
            'count': summary['pages_without_schema']
        })

    if summary['missing_canonical'] > 0 and len(priorities) < 3:
        pct = round((summary['missing_canonical'] / total_pages) * 100, 1)
        priorities.append({
            'text': f"🟠 Dodaj Canonical do {summary['missing_canonical']}/{total_pages} stron ({pct}%)",
            'desc': "Zapobiega problemom z duplikacją treści",
            'count': summary['missing_canonical']
        })

    if summary['pages_without_viewport'] > 0 and len(priorities) < 3:
        pct = round((summary['pages_without_viewport'] / total_pages) * 100, 1)
        priorities.append({
            'text': f"🟡 Dodaj meta viewport do {summary['pages_without_viewport']}/{total_pages} stron ({pct}%)",
            'desc': "Kluczowe dla mobile-first indexing Google",
            'count': summary['pages_without_viewport']
        })

    # Sortujemy po count (malejąco) i bierzemy top 3
    priorities.sort(key=lambda x: x['count'], reverse=True)

    for i, pr in enumerate(priorities[:3], 1):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{i}. {pr['text']}\n")
        run.bold = True
        run2 = p.add_run(f"   {pr['desc']}")
        run2.font.size = Pt(10)
        run2.italic = True
        run2.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    stats_table = doc.add_table(rows=1, cols=3)
    stats_table.style = 'Light Grid Accent 1'
    hdr = stats_table.rows[0].cells
    hdr[0].text = 'Metryka'
    hdr[1].text = 'Wartość'
    hdr[2].text = 'Opis'

    # Funkcja pomocnicza do dodawania wierszy z opisem
    def add_metric_row(metric_name, value, description=""):
        row = stats_table.add_row().cells
        row[0].text = metric_name
        row[1].text = str(value)
        row[2].text = description

    # Metryki z opisami
    add_metric_row(
        'Przeskanowanych stron',
        f"{summary['pages_crawled']}",
        "Całkowita liczba odkrytych URL-i podczas crawlingu"
    )

    add_metric_row(
        'Przeanalizowanych stron',
        f"{summary['pages_analyzed']} (wykluczono {summary['pages_excluded']})",
        "Strony poddane analizie SEO (bez technicznych URL jak /cdn-cgi/*)"
    )

    add_metric_row(
        'Strony OK (200)',
        f"✅ {summary['pages_ok']}",
        "Strony działające poprawnie, zwracające kod HTTP 200"
    )

    add_metric_row(
        'Strony z błędami (4xx)',
        f"🔴 {summary['pages_with_errors']}",
        "Błędy klienta (404 Not Found, 403 Forbidden itp.) - blokują indeksację"
    )

    add_metric_row(
        'Brak Title',
        f"{summary['missing_title']}",
        "Strony bez tagu <title> - kluczowego dla pozycjonowania i CTR"
    )

    add_metric_row(
        'Brak Description',
        f"{summary['missing_description']}",
        "Strony bez <meta name=\"description\"> - wpływa na snippet w Google"
    )

    add_metric_row(
        'Problemy Title (długość)',
        f"{summary['title_issues']}",
        "Title zbyt krótki (<30 znaków) lub za długi (>65) - Google może obciąć"
    )

    add_metric_row(
        'Problemy Description (długość)',
        f"{summary['description_issues']}",
        "Description poza zakresem 120-165 znaków - może być obcięty lub zbyt krótki"
    )

    add_metric_row(
        'Duplikaty Title',
        f"{summary['duplicate_titles']}",
        "Różne strony z identycznym tytułem - konfunduje Google i użytkowników"
    )

    add_metric_row(
        'Duplikaty Description',
        f"{summary['duplicate_descriptions']}",
        "Różne strony z tą samą meta description - obniża unikatowość"
    )

    add_metric_row(
        'Brak Canonical',
        f"{summary['missing_canonical']}",
        "Brak tagu <link rel=\"canonical\"> - prowadzi do problemów z duplikacją treści"
    )

    add_metric_row(
        'Brak H1',
        f"{summary['missing_h1']}",
        "Strona bez nagłówka głównego <h1> - kluczowy element struktury i SEO"
    )

    add_metric_row(
        'Wiele H1',
        f"{summary['multiple_h1']}",
        "Strona z >1 nagłówkiem H1 - może wprowadzać wyszukiwarki w błąd"
    )

    add_metric_row(
        'Obrazy bez ALT (łączna liczba)',
        f"{summary['total_images_without_alt']}",
        "Obrazki bez atrybutu alt - problem dla SEO, dostępności i obrazków Google"
    )

    add_metric_row(
        'Mobile-friendly',
        f"{summary['mobile_friendly_pages']} ({summary['mobile_percentage']}%)",
        "Strony dostosowane do urządzeń mobilnych (responsive design)"
    )

    add_metric_row(
        'Brak meta viewport',
        f"{summary['pages_without_viewport']}",
        "Strony bez <meta name=\"viewport\"> - nie skalują się na mobile"
    )

    add_metric_row(
        'Brak Open Graph',
        f"{summary['pages_without_og']}",
        "Strony bez meta tagów OG (Facebook, LinkedIn) - brzydki podgląd linków"
    )

    add_metric_row(
        'Brak Twitter Cards',
        f"{summary['pages_without_twitter']}",
        "Strony bez Twitter Card - brak atrakcyjnego podglądu na X (Twitter)"
    )

    add_metric_row(
        'Strony z danymi strukturalnymi',
        f"{summary['pages_with_schema']}",
        "Liczba stron z Schema.org (JSON-LD) - umożliwia rich snippets w Google"
    )

    add_metric_row(
        'Brak Schema',
        f"{summary['pages_without_schema']}",
        "Strony bez Schema.org - tracisz gwiazdki, FAQ, breadcrumbs w wynikach"
    )

    add_metric_row(
        'Śr. typów Schema/stronę',
        f"{summary['avg_schema_types']}",
        "Ile różnych typów Schema jest na stronie (Article, Product, FAQ, etc.)"
    )

    add_metric_row(
        'Śr. E-E-A-T',
        f"{summary['avg_eeat_score']}%",
        "Experience, Expertise, Authoritativeness, Trustworthiness - sygnały jakości Google"
    )

    add_metric_row(
        'Słabe E-E-A-T',
        f"{summary['pages_weak_eeat']}",
        "Strony z niskim E-E-A-T (<50%): brak autora, dat, certyfikatów, źródeł"
    )

    add_metric_row(
        'Local NAP OK',
        f"{summary['local_optimized_pages']}",
        "Strony z poprawnymi danymi NAP (Name, Address, Phone) - istotne dla firm lokalnych"
    )

    add_metric_row(
        'Słaby Local SEO',
        f"{summary['pages_poor_local_seo']}",
        "Strony bez NAP, Schema LocalBusiness, linków do mapy - słabo dla SEO lokalnego"
    )

    add_metric_row(
        'Thin content (<300 słów)',
        f"{summary['thin_content_pages']}",
        "Strony z bardzo krótką treścią - Google może uznać za low-quality"
    )

    add_metric_row(
        'Śr. Security',
        f"{summary['avg_security_score']}%",
        "Średni poziom zabezpieczeń (HTTPS + security headers). Tu: {:.1f}% to bardzo nisko".format(
            summary['avg_security_score'])
    )

    add_metric_row(
        'Słabe bezpieczeństwo',
        f"{summary['pages_poor_security']}",
        "Strony z oceną <50%: brak kluczowych nagłówków (HSTS, CSP, X-Frame-Options)"
    )

    add_metric_row(
        'Brakujące security headers',
        f"{summary['pages_missing_security_headers']}",
        "Strony z <3 nagłówkami security. Sprawdź: HSTS, CSP, X-Frame-Options, X-Content-Type-Options, Referrer-Policy"
    )

    add_metric_row(
        'Mixed content',
        f"{summary['pages_with_mixed_content']}",
        "Strony HTTPS z zasobami HTTP (obrazki, skrypty) - warning w przeglądarce"
    )

    add_metric_row(
        'WYNIK AUDYTU',
        f"{summary['overall_score']}/100 ({summary['overall_grade']})",
        "Ocena łączna: Availability (30%) + Meta (15%) + Mobile (15%) + Schema (10%) + E-E-A-T (10%) + Security (20%)"
    )
    doc.add_page_break()

    doc.add_heading('2. 🎯 Priorytety (od krytycznych do lekkich)', 1)

    if issues['critical_errors']:
        doc.add_heading('🔴 Błędy HTTP (4xx)', 2)
        p = doc.add_paragraph();
        p.add_run(f"Znaleziono {len(issues['critical_errors'])} stron z błędami 4xx.").bold = True
        doc.add_paragraph(
            "Strony te są niedostępne dla użytkowników i robotów Google, co skutkuje:\n• Utratą ruchu organicznego\n• Negatywnym wpływem na UX\n• Problemami z indeksacją")
        doc.add_paragraph()
        p = doc.add_paragraph();
        run = p.add_run("Problematyczne adresy:");
        run.bold = True
        for err in issues['critical_errors'][:20]:
            status_code = err.get('status', 'N/A')
            error_msg = err.get('error', '')
            if error_msg and error_msg != '':
                doc.add_paragraph(f"• {err['url']} - Status: {status_code} ({error_msg})", style='List Bullet')
            else:
                doc.add_paragraph(f"• {err['url']} - Status: {status_code}", style='List Bullet')
        if len(issues['critical_errors']) > 20:
            doc.add_paragraph(f"...oraz {len(issues['critical_errors']) - 20} innych")

    if issues['missing_title'] or issues['missing_description'] or issues['missing_canonical']:
        doc.add_heading('🟠 Meta & Canonical', 2)

        if issues['missing_title']:
            p = doc.add_paragraph();
            run = p.add_run(f"Brak Title: {len(issues['missing_title'])} stron");
            run.bold = True
            doc.add_paragraph(
                "Tag <title> to pierwszy element, jaki użytkownik widzi w wynikach Google. Jego brak oznacza:")
            doc.add_paragraph("• Brak kontroli nad tym, co Google wyświetli w SERP", style='List Bullet')
            doc.add_paragraph("• Niższy CTR (Click-Through Rate)", style='List Bullet')
            doc.add_paragraph("• Słabsze pozycjonowanie", style='List Bullet')
            doc.add_paragraph()
            p = doc.add_paragraph();
            run = p.add_run("Strony bez Title:");
            run.bold = True
            for url in issues['missing_title'][:15]:
                doc.add_paragraph(f"• {url}", style='List Bullet')
            if len(issues['missing_title']) > 15:
                doc.add_paragraph(f"...oraz {len(issues['missing_title']) - 15} innych")
            doc.add_paragraph()

        if issues['missing_description']:
            p = doc.add_paragraph();
            run = p.add_run(f"Brak Meta Description: {len(issues['missing_description'])} stron");
            run.bold = True
            doc.add_paragraph("Meta description to 'zachęta' do kliknięcia w wynikach wyszukiwania. Bez niej:")
            doc.add_paragraph("• Google sam generuje opis (często nietrafiający w sedno)", style='List Bullet')
            doc.add_paragraph("• Tracisz kontrolę nad przekazem marketingowym", style='List Bullet')
            doc.add_paragraph("• CTR może spaść nawet o 30-40%", style='List Bullet')
            doc.add_paragraph()
            p = doc.add_paragraph();
            run = p.add_run("Strony bez Meta Description:");
            run.bold = True
            for url in issues['missing_description'][:15]:
                doc.add_paragraph(f"• {url}", style='List Bullet')
            if len(issues['missing_description']) > 15:
                doc.add_paragraph(f"...oraz {len(issues['missing_description']) - 15} innych")
            doc.add_paragraph()

        if issues['missing_canonical']:
            p = doc.add_paragraph();
            run = p.add_run(f"Brak Canonical: {len(issues['missing_canonical'])} stron");
            run.bold = True
            doc.add_paragraph(
                "Tag canonical wskazuje Google, która wersja strony jest 'główna'. Jego brak prowadzi do:")
            doc.add_paragraph("• Problemów z duplikacją treści (Google nie wie, którą wersję indeksować)",
                              style='List Bullet')
            doc.add_paragraph("• Rozproszenia 'mocy' linków między duplikatami", style='List Bullet')
            doc.add_paragraph("• Słabszego pozycjonowania wszystkich wersji", style='List Bullet')
            doc.add_paragraph()
            p = doc.add_paragraph();
            run = p.add_run("Strony bez Canonical:");
            run.bold = True
            for url in issues['missing_canonical'][:15]:
                doc.add_paragraph(f"• {url}", style='List Bullet')
            if len(issues['missing_canonical']) > 15:
                doc.add_paragraph(f"...oraz {len(issues['missing_canonical']) - 15} innych")
            doc.add_paragraph()

    if issues['poor_security'] or issues['missing_security_headers']:
        doc.add_heading('🟠 Bezpieczeństwo – niski poziom / brak nagłówków', 2)

        if issues['poor_security']:
            p = doc.add_paragraph();
            run = p.add_run(f"Słabe bezpieczeństwo: {len(issues['poor_security'])} stron");
            run.bold = True
            doc.add_paragraph(
                "Strony z oceną bezpieczeństwa <50% mają braki w podstawowych nagłówkach zabezpieczających:")
            doc.add_paragraph()
            p = doc.add_paragraph();
            run = p.add_run("Przykłady stron ze słabym security:");
            run.bold = True
            for item in issues['poor_security'][:10]:
                url = item['url']
                sec_pct = item.get('security_percentage', 0)
                missing = item.get('missing_headers', [])
                if missing:
                    missing_str = ", ".join(missing[:3])
                    doc.add_paragraph(f"• {url} ({sec_pct}%) - Brak: {missing_str}", style='List Bullet')
                else:
                    doc.add_paragraph(f"• {url} ({sec_pct}%)", style='List Bullet')
            if len(issues['poor_security']) > 10:
                doc.add_paragraph(f"...oraz {len(issues['poor_security']) - 10} innych")
            doc.add_paragraph()

        if issues['missing_security_headers']:
            p = doc.add_paragraph();
            run = p.add_run(f"Braki w security headers: {len(issues['missing_security_headers'])} stron");
            run.bold = True
            doc.add_paragraph(
                "Strony z mniej niż 3 nagłówkami bezpieczeństwa są podatne na ataki. Brak odpowiednich headerów oznacza:")
            doc.add_paragraph("• Łatwiejsze przeprowadzenie ataków XSS, clickjacking", style='List Bullet')
            doc.add_paragraph("• Brak wymuszenia HTTPS (możliwy man-in-the-middle)", style='List Bullet')
            doc.add_paragraph("• Obniżone zaufanie użytkowników i Google", style='List Bullet')
            doc.add_paragraph()
            p = doc.add_paragraph();
            run = p.add_run("Kluczowe brakujące nagłówki:");
            run.bold = True
            p = doc.add_paragraph();
            run = p.add_run("• HSTS: ");
            run.bold = True
            run2 = p.add_run("Wymusza HTTPS, chroni przed atakami man-in-the-middle")
            p = doc.add_paragraph();
            run = p.add_run("• CSP: ");
            run.bold = True
            run2 = p.add_run("Zapobiega atakom XSS (wstrzykiwanie złośliwego kodu)")
            p = doc.add_paragraph();
            run = p.add_run("• X-Frame-Options: ");
            run.bold = True
            run2 = p.add_run("Chroni przed clickjacking (osadzenie strony w iframe)")
            doc.add_paragraph()

    if issues['no_viewport'] or issues['no_og_tags'] or issues['no_twitter_cards'] or issues['missing_schema']:
        doc.add_heading('🟡 Mobile / Social / Schema', 2)

        if issues['no_viewport']:
            p = doc.add_paragraph();
            run = p.add_run(f"Brak meta viewport: {len(issues['no_viewport'])} stron");
            run.bold = True
            doc.add_paragraph("Od 2018 Google stosuje mobile-first indexing. Brak meta viewport oznacza, że:")
            doc.add_paragraph("• Strona nie skaluje się poprawnie na smartfonach", style='List Bullet')
            doc.add_paragraph("• Google może obniżyć ranking (mobile-first!)", style='List Bullet')
            doc.add_paragraph("• Użytkownicy mobile widzą 'desktopową' wersję (zła UX)", style='List Bullet')
            doc.add_paragraph()

        if issues['no_og_tags']:
            p = doc.add_paragraph();
            run = p.add_run(f"Braki w Open Graph: {len(issues['no_og_tags'])} stron");
            run.bold = True
            doc.add_paragraph(
                "Open Graph to meta tagi używane przez Facebook, LinkedIn, WhatsApp do generowania podglądu linków.")
            doc.add_paragraph()

        if issues['no_twitter_cards']:
            p = doc.add_paragraph();
            run = p.add_run(f"Brak Twitter Cards: {len(issues['no_twitter_cards'])} stron");
            run.bold = True
            doc.add_paragraph("Twitter Cards to odpowiednik OG dla platformy X (dawniej Twitter).")
            doc.add_paragraph()

        if issues['missing_schema']:
            p = doc.add_paragraph();
            run = p.add_run(f"Brak Schema.org: {len(issues['missing_schema'])} stron");
            run.bold = True
            doc.add_paragraph(
                "Schema.org (JSON-LD) to 'język', którym mówisz do Google o zawartości strony. Bez niego:")
            doc.add_paragraph("• Tracisz rich snippets (gwiazdki, FAQ, breadcrumbs)", style='List Bullet')
            doc.add_paragraph("• Trudniej o featured snippet (pozycja 0)", style='List Bullet')
            doc.add_paragraph("• Google słabiej rozumie kontekst treści", style='List Bullet')
            doc.add_paragraph()

    if issues['weak_eeat'] or issues['thin_content']:
        doc.add_heading('🟡 E-E-A-T & Treść', 2)

        if issues['weak_eeat']:
            p = doc.add_paragraph();
            run = p.add_run(f"Słabe E-E-A-T: {len(issues['weak_eeat'])} stron");
            run.bold = True
            doc.add_paragraph(
                "E-E-A-T (Experience, Expertise, Authoritativeness, Trustworthiness) to zestaw sygnałów jakości dla Google.")
            doc.add_paragraph()

        if issues['thin_content']:
            p = doc.add_paragraph();
            run = p.add_run(f"Thin content: {len(issues['thin_content'])} stron");
            run.bold = True
            doc.add_paragraph("Strony z mniej niż 300 słowami mogą być uznane przez Google za 'cienkie' (low-quality).")
            doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading('5. 📱 Mobilność i Responsywność', 1)
    p = doc.add_paragraph();
    run = p.add_run(f'Status mobilności: ');
    run.bold = True
    if summary['mobile_percentage'] >= 90:
        run = p.add_run(f"✅ {summary['mobile_percentage']}% stron mobile-friendly");
        run.font.color.rgb = RGBColor(0, 150, 0)
    elif summary['mobile_percentage'] >= 70:
        run = p.add_run(f"⚠️ {summary['mobile_percentage']}% stron mobile-friendly");
        run.font.color.rgb = RGBColor(200, 100, 0)
    else:
        run = p.add_run(f"🔴 {summary['mobile_percentage']}% stron mobile-friendly");
        run.font.color.rgb = RGBColor(200, 0, 0)

    doc.add_paragraph()
    doc.add_paragraph(
        "Od marca 2018 Google stosuje mobile-first indexing - najpierw analizuje wersję mobilną strony. Brak responsywności oznacza:")
    doc.add_paragraph("• Spadek pozycji w wynikach mobilnych (60%+ ruchu to mobile)", style='List Bullet')
    doc.add_paragraph("• Gorsze doświadczenie użytkownika = wyższy bounce rate", style='List Bullet')
    doc.add_paragraph("• Utratę potencjalnych klientów mobilnych", style='List Bullet')

    if issues['no_viewport']:
        doc.add_paragraph()
        p = doc.add_paragraph();
        run = p.add_run(f"Strony bez meta viewport ({len(issues['no_viewport'])}):");
        run.bold = True
        doc.add_paragraph(
            "Te strony nie dostosują się do rozmiaru ekranu smartfona - użytkownik zobaczy miniaturkę wersji desktopowej.")
        doc.add_paragraph()

        # Filtrujemy pliki multimedialne (mp4, mp3, jpg, png, pdf, etc.)
        multimedia_extensions = ['.mp4', '.mp3', '.avi', '.mov', '.wmv', '.flv', '.webm', '.mkv',
                                 '.jpg', '.jpeg', '.png', '.gif', '.webp', '.svg', '.pdf', '.zip',
                                 '.rar', '.doc', '.docx', '.xls', '.xlsx']

        filtered_urls = []
        for url in issues['no_viewport']:
            url_lower = url.lower()
            is_multimedia = any(url_lower.endswith(ext) for ext in multimedia_extensions)
            if not is_multimedia:
                filtered_urls.append(url)

        if filtered_urls:
            p = doc.add_paragraph();
            run = p.add_run("Strony HTML wymagające poprawy:");
            run.bold = True;
            run.font.size = Pt(11)
            for url in filtered_urls[:15]:
                doc.add_paragraph(f"• {url}", style='List Bullet')
            if len(filtered_urls) > 15:
                doc.add_paragraph(f"...oraz {len(filtered_urls) - 15} innych stron")

        # Jeśli są pliki multimedialne, informujemy o nich osobno
        multimedia_count = len(issues['no_viewport']) - len(filtered_urls)
        if multimedia_count > 0:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(f"ℹ️  Dodatkowo znaleziono {multimedia_count} plików multimedialnych bez viewport ")
            run.font.size = Pt(9)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            run2 = p.add_run("(pliki video/obrazy - to normalne, nie wymaga poprawy)")
            run2.font.size = Pt(9)
            run2.italic = True
            run2.font.color.rgb = RGBColor(100, 100, 100)

        if SHOW_REMEDIATIONS:
            doc.add_paragraph()
            p = doc.add_paragraph();
            run = p.add_run('Jak dodać meta viewport:\n');
            run.bold = True
            code = '<meta name="viewport" content="width=device-width, initial-scale=1.0">'
            q = doc.add_paragraph();
            r = q.add_run(code);
            r.font.name = 'Courier New';
            r.font.size = Pt(10)

    doc.add_page_break()
    doc.add_heading('6. 📢 Open Graph i Twitter Cards', 1)

    doc.add_paragraph(
        "Social media to potężne źródło ruchu. Gdy ktoś udostępnia link do Twojej strony na Facebooku, LinkedIn czy X (Twitter), te platformy generują 'podgląd' - miniaturkę z obrazkiem, tytułem i opisem. To pierwsze wrażenie decyduje, czy użytkownik kliknie.")
    doc.add_paragraph()

    if issues['no_og_tags']:
        p = doc.add_paragraph();
        run = p.add_run(f"Braki w Open Graph: {len(issues['no_og_tags'])} stron");
        run.bold = True
        doc.add_paragraph("Open Graph to meta tagi używane przez Facebook, LinkedIn, WhatsApp, Messenger. Bez nich:")
        doc.add_paragraph("• Platformy same wybierają obrazek (często nietrafiający w sedno)", style='List Bullet')
        doc.add_paragraph("• Opis może być losowy fragment tekstu", style='List Bullet')
        doc.add_paragraph("• Tracisz kontrolę nad przekazem marketingowym", style='List Bullet')
        doc.add_paragraph("• Niższy CTR z social media (nawet o 50%!)", style='List Bullet')
        doc.add_paragraph()

        p = doc.add_paragraph();
        run = p.add_run("Strony wymagające dodania Open Graph:");
        run.bold = True;
        run.font.size = Pt(11)
        for item in issues['no_og_tags'][:15]:
            url = item['url']
            has_image = item.get('has_og_image', False)
            has_title = item.get('has_og_title', False)
            has_desc = item.get('has_og_description', False)

            missing_parts = []
            if not has_image: missing_parts.append("obrazek")
            if not has_title: missing_parts.append("tytuł")
            if not has_desc: missing_parts.append("opis")

            if missing_parts:
                doc.add_paragraph(f"• {url}\n  Brak: {', '.join(missing_parts)}", style='List Bullet')
            else:
                doc.add_paragraph(f"• {url}", style='List Bullet')

        if len(issues['no_og_tags']) > 15:
            doc.add_paragraph(f"...oraz {len(issues['no_og_tags']) - 15} innych stron")
        doc.add_paragraph()

    if issues['no_twitter_cards']:
        p = doc.add_paragraph();
        run = p.add_run(f"Brak Twitter Cards: {len(issues['no_twitter_cards'])} stron");
        run.bold = True
        doc.add_paragraph(
            "Twitter Cards to odpowiednik Open Graph dla platformy X (dawniej Twitter). Działają analogicznie - kontrolują, jak Twój link wygląda po udostępnieniu.")
        doc.add_paragraph()

        p = doc.add_paragraph();
        run = p.add_run("Strony wymagające dodania Twitter Cards:");
        run.bold = True;
        run.font.size = Pt(11)
        for url in issues['no_twitter_cards'][:15]:
            doc.add_paragraph(f"• {url}", style='List Bullet')
        if len(issues['no_twitter_cards']) > 15:
            doc.add_paragraph(f"...oraz {len(issues['no_twitter_cards']) - 15} innych stron")
        doc.add_paragraph()

    if SHOW_REMEDIATIONS:
        example_code = '''<!-- Open Graph -->
<meta property="og:title" content="Tytuł" />
<meta property="og:description" content="Opis" />
<meta property="og:image" content="https://twoja-domena.pl/obraz.jpg" />
<meta property="og:url" content="https://twoja-domena.pl/strona" />
<meta property="og:type" content="website" />
<!-- Twitter Cards -->
<meta name="twitter:card" content="summary_large_image" />
<meta name="twitter:title" content="Tytuł" />
<meta name="twitter:description" content="Opis" />
<meta name="twitter:image" content="https://twoja-domena.pl/obraz.jpg" />'''
        p = doc.add_paragraph();
        run = p.add_run('Przykładowe tagi:');
        run.bold = True
        q = doc.add_paragraph();
        r = q.add_run(example_code);
        r.font.name = 'Courier New';
        r.font.size = Pt(9)

    doc.add_page_break()
    doc.add_heading('7. 🔗 Dane Strukturalne (Schema.org)', 1)

    doc.add_paragraph(
        "Schema.org (JSON-LD) to 'język techniczny', którym mówisz Google o zawartości swojej strony. To nie magia, ale konkretne instrukcje: 'to jest artykuł', 'to jest produkt z ceną', 'to FAQ z pytaniami i odpowiedziami'. Dzięki temu Google może wyświetlić Twoją stronę w bardziej atrakcyjny sposób w wynikach wyszukiwania - tzw. rich snippets.")
    doc.add_paragraph()

    schema_percentage = (summary['pages_with_schema'] / max(1, summary['pages_analyzed'])) * 100

    if schema_percentage >= 70:
        p = doc.add_paragraph();
        run = p.add_run(f"✅ {summary['pages_with_schema']} stron ma dane strukturalne ({schema_percentage:.1f}%)");
        run.font.color.rgb = RGBColor(0, 150, 0);
        run.bold = True
    else:
        p = doc.add_paragraph();
        run = p.add_run(
            f"⚠️ Tylko {summary['pages_with_schema']} stron ma dane strukturalne ({schema_percentage:.1f}%)");
        run.font.color.rgb = RGBColor(200, 100, 0);
        run.bold = True

    doc.add_paragraph()
    doc.add_paragraph("Co tracisz bez Schema.org:")
    doc.add_paragraph("• Rich snippets: gwiazdki ocen, ceny produktów, FAQ rozwijane w SERP", style='List Bullet')
    doc.add_paragraph("• Breadcrumbs (ścieżka nawigacji) w wynikach Google", style='List Bullet')
    doc.add_paragraph("• Featured snippet (pozycja 0) - trudniej bez struktury", style='List Bullet')
    doc.add_paragraph("• Karuzele produktów/artykułów w mobilnych wynikach", style='List Bullet')
    doc.add_paragraph("• Lepsze zrozumienie kontekstu przez Google (ważne dla AI)", style='List Bullet')
    doc.add_paragraph()

    if issues['missing_schema']:
        p = doc.add_paragraph();
        run = p.add_run(f"Strony bez Schema.org: {len(issues['missing_schema'])}");
        run.bold = True
        doc.add_paragraph()

        p = doc.add_paragraph();
        run = p.add_run("Przykłady stron wymagających Schema:");
        run.bold = True;
        run.font.size = Pt(11)
        for url in issues['missing_schema'][:20]:
            doc.add_paragraph(f"• {url}", style='List Bullet')
        if len(issues['missing_schema']) > 20:
            doc.add_paragraph(f"...oraz {len(issues['missing_schema']) - 20} innych stron")

        doc.add_paragraph()
        p = doc.add_paragraph();
        run = p.add_run("💡 Rekomendacja:");
        run.bold = True;
        run.font.color.rgb = RGBColor(50, 100, 200)
        doc.add_paragraph("Priorytetowo dodaj Schema do:")
        doc.add_paragraph("1. Strony głównej (Organization/LocalBusiness)", style='List Number')
        doc.add_paragraph("2. Stron produktów (Product z ceną i dostępnością)", style='List Number')
        doc.add_paragraph("3. Artykułów blogowych (Article/BlogPosting)", style='List Number')
        doc.add_paragraph("4. FAQ/Pytania (FAQPage)", style='List Number')
        doc.add_paragraph("5. Opinii klientów (Review/AggregateRating)", style='List Number')

    doc.add_page_break()
    doc.add_heading('8. 🏆 E-E-A-T', 1)

    doc.add_paragraph(
        "E-E-A-T to akronim od Experience, Expertise, Authoritativeness, Trustworthiness - po polsku: Doświadczenie, Eksperckość, Autorytet, Zaufanie. To nie jest 'ranking factor' w klasycznym sensie, ale zestaw sygnałów, które Quality Raterzy Google (ludzie oceniający jakość wyników) sprawdzają ręcznie. Algorytm uczy się na ich ocenach.")
    doc.add_paragraph()

    doc.add_paragraph("Co sprawdza Google oceniając E-E-A-T:")
    doc.add_paragraph("• Czy autor jest ekspertem w temacie? (widoczne imię, nazwisko, bio)", style='List Bullet')
    doc.add_paragraph("• Czy treść jest aktualna? (data publikacji, data aktualizacji)", style='List Bullet')
    doc.add_paragraph("• Czy strona jest autorytetem? (linki z wiarygodnych źródeł .edu/.gov)", style='List Bullet')
    doc.add_paragraph("• Czy można zaufać? (HTTPS, dane kontaktowe, polityka prywatności)", style='List Bullet')
    doc.add_paragraph("• Czy są opinie/recenzje? (social proof)", style='List Bullet')
    doc.add_paragraph()

    p = doc.add_paragraph();
    run = p.add_run(f'Średnia ocena E-E-A-T: ');
    run.bold = True
    if summary['avg_eeat_score'] >= 70:
        run = p.add_run(f"✅ {summary['avg_eeat_score']}%");
        run.font.color.rgb = RGBColor(0, 150, 0)
    elif summary['avg_eeat_score'] >= 50:
        run = p.add_run(f"⚠️ {summary['avg_eeat_score']}%");
        run.font.color.rgb = RGBColor(200, 100, 0)
    else:
        run = p.add_run(f"🔴 {summary['avg_eeat_score']}%");
        run.font.color.rgb = RGBColor(200, 0, 0)

    doc.add_paragraph()

    if issues['weak_eeat']:
        p = doc.add_paragraph();
        run = p.add_run(f"Strony ze słabym E-E-A-T (<50%): {len(issues['weak_eeat'])}");
        run.bold = True
        doc.add_paragraph(
            "Te strony mają niską ocenę zaufania w oczach Google. To nie znaczy, że są 'złe' - po prostu brakuje im sygnałów jakości.")
        doc.add_paragraph()

        p = doc.add_paragraph();
        run = p.add_run("Przykłady stron ze słabym E-E-A-T:");
        run.bold = True;
        run.font.size = Pt(11)
        for item in issues['weak_eeat'][:15]:
            url = item['url']
            eeat_pct = item.get('eeat_percentage', 0)
            missing = item.get('missing', [])

            missing_readable = []
            for key in missing:
                if key == 'has_author':
                    missing_readable.append("autor")
                elif key == 'has_date':
                    missing_readable.append("data publikacji")
                elif key == 'has_expertise_signals':
                    missing_readable.append("sygnały eksperckości")
                elif key == 'has_quality_external_links':
                    missing_readable.append("linki do źródeł")
                elif key == 'has_contact_info':
                    missing_readable.append("dane kontaktowe")
                elif key == 'has_reviews':
                    missing_readable.append("opinie")

            if missing_readable:
                doc.add_paragraph(f"• {url} ({eeat_pct:.1f}%)\n  Brak: {', '.join(missing_readable[:3])}",
                                  style='List Bullet')
            else:
                doc.add_paragraph(f"• {url} ({eeat_pct:.1f}%)", style='List Bullet')

        if len(issues['weak_eeat']) > 15:
            doc.add_paragraph(f"...oraz {len(issues['weak_eeat']) - 15} innych stron")

        doc.add_paragraph()
        p = doc.add_paragraph();
        run = p.add_run("💡 Jak poprawić E-E-A-T:");
        run.bold = True;
        run.font.color.rgb = RGBColor(50, 100, 200)
        doc.add_paragraph("1. Dodaj ramkę 'O autorze' z imieniem, nazwiskiem, zdjęciem, bio", style='List Number')
        doc.add_paragraph("2. Wstaw datę publikacji i 'Ostatnia aktualizacja: [data]'", style='List Number')
        doc.add_paragraph("3. Linkuj do wiarygodnych źródeł (.edu, .gov, badania naukowe)", style='List Number')
        doc.add_paragraph("4. Dodaj certyfikaty, nagrody, doświadczenie firmy", style='List Number')
        doc.add_paragraph("5. Umieść widoczne dane kontaktowe (telefon, email, adres)", style='List Number')

    doc.add_page_break()
    doc.add_heading('9. 📍 Local SEO (NAP)', 1)

    doc.add_paragraph(
        "NAP to skrót od Name, Address, Phone - czyli nazwa firmy, adres i telefon. To podstawa lokalnego SEO. Jeśli prowadzisz biznes stacjonarny (sklep, biuro, restauracja, salon) lub obsługujesz określony region, Google sprawdza, czy Twoje dane NAP są:")
    doc.add_paragraph("• Spójne (takie same wszędzie: strona, Google Maps, Facebook, wizytówki)", style='List Bullet')
    doc.add_paragraph("• Widoczne (łatwo znaleźć na stronie)", style='List Bullet')
    doc.add_paragraph("• Ustrukturyzowane (Schema.org LocalBusiness)", style='List Bullet')
    doc.add_paragraph()

    local_percentage = (summary['local_optimized_pages'] / max(1, summary['pages_analyzed'])) * 100

    if local_percentage >= 50:
        p = doc.add_paragraph();
        run = p.add_run(
            f"✅ {summary['local_optimized_pages']} stron zoptymalizowanych pod NAP ({local_percentage:.1f}%)");
        run.font.color.rgb = RGBColor(0, 150, 0);
        run.bold = True
    else:
        p = doc.add_paragraph();
        run = p.add_run(f"⚠️ Tylko {summary['local_optimized_pages']} stron posiada NAP ({local_percentage:.1f}%)");
        run.font.color.rgb = RGBColor(200, 100, 0);
        run.bold = True

    doc.add_paragraph()
    doc.add_paragraph("Dlaczego NAP jest ważny:")
    doc.add_paragraph("• Google Local Pack (3 wyniki na mapie) wymaga spójnych danych", style='List Bullet')
    doc.add_paragraph("• Użytkownicy szukający 'firma + miasto' trafiają na lokalne wyniki", style='List Bullet')
    doc.add_paragraph("• Zaufanie: widoczny telefon i adres = większa konwersja", style='List Bullet')
    doc.add_paragraph("• Voice search ('Hey Google, znajdź X w pobliżu') preferuje NAP", style='List Bullet')
    doc.add_paragraph()

    if issues['poor_local_seo']:
        p = doc.add_paragraph();
        run = p.add_run(f"Strony bez poprawnego NAP: {len(issues['poor_local_seo'])}");
        run.bold = True
        doc.add_paragraph()

        p = doc.add_paragraph();
        run = p.add_run("Przykłady stron wymagających poprawy Local SEO:");
        run.bold = True;
        run.font.size = Pt(11)
        for item in issues['poor_local_seo'][:15]:
            url = item['url']
            nap_score = item.get('nap_score', 0)
            phones = item.get('phone_numbers', 0)
            has_address = item.get('has_address', False)
            has_schema = item.get('has_local_schema', False)

            issues_list = []
            if phones == 0: issues_list.append("brak telefonu")
            if not has_address: issues_list.append("brak adresu")
            if not has_schema: issues_list.append("brak Schema LocalBusiness")

            if issues_list:
                doc.add_paragraph(f"• {url} (NAP: {nap_score}/3)\n  Problem: {', '.join(issues_list)}",
                                  style='List Bullet')
            else:
                doc.add_paragraph(f"• {url} (NAP: {nap_score}/3)", style='List Bullet')

        if len(issues['poor_local_seo']) > 15:
            doc.add_paragraph(f"...oraz {len(issues['poor_local_seo']) - 15} innych stron")

        doc.add_paragraph()
        p = doc.add_paragraph();
        run = p.add_run("💡 Jak poprawić Local SEO:");
        run.bold = True;
        run.font.color.rgb = RGBColor(50, 100, 200)
        doc.add_paragraph("1. Dodaj widoczną stopkę z: nazwa firmy, adres, telefon, email", style='List Number')
        doc.add_paragraph("2. Wdróż Schema.org LocalBusiness (JSON-LD) z pełnymi danymi NAP", style='List Number')
        doc.add_paragraph("3. Sprawdź spójność: te same dane na Google Maps, Facebook, stronie", style='List Number')
        doc.add_paragraph("4. Dodaj mapę Google (embed) ze wskazaniem lokalizacji", style='List Number')
        doc.add_paragraph("5. Stwórz dedykowaną podstronę 'Kontakt' z pełnymi danymi", style='List Number')

    doc.add_page_break()
    doc.add_heading('10. 📝 Jakość Treści', 1)

    doc.add_paragraph(
        "Google nie lubi 'cienkich' stron - czyli takich, które mają bardzo mało tekstu (poniżej 300 słów). Dlaczego? Bo algorytm zakłada, że krótka strona = mało wartościowa informacja = niska jakość. Oczywiście są wyjątki (strona kontaktu, landing page produktowy), ale generalnie: im więcej merytorycznej treści, tym lepiej.")
    doc.add_paragraph()

    doc.add_paragraph("Co ryzykujesz mając 'thin content':")
    doc.add_paragraph("• Google może uznać stronę za low-quality i obniżyć jej ranking", style='List Bullet')
    doc.add_paragraph("• Trudniej o featured snippet (pozycja 0) - potrzeba więcej kontekstu", style='List Bullet')
    doc.add_paragraph("• Użytkownicy szybko opuszczają stronę (wysoki bounce rate)", style='List Bullet')
    doc.add_paragraph("• Mniejsza szansa na linki zewnętrzne (kto zlinkuje do 100-słownego tekstu?)",
                      style='List Bullet')
    doc.add_paragraph()

    if issues['thin_content']:
        p = doc.add_paragraph();
        run = p.add_run(f"Strony z thin content (<300 słów): {len(issues['thin_content'])}");
        run.bold = True
        doc.add_paragraph()

        p = doc.add_paragraph();
        run = p.add_run("Strony wymagające rozbudowania treści:");
        run.bold = True;
        run.font.size = Pt(11)
        for item in issues['thin_content']:
            url = item['url']
            word_count = item.get('word_count', 0)
            text_len = item.get('text_len', 0)

            doc.add_paragraph(f"• {url}\n  Słów: {word_count}, Znaków: {text_len}", style='List Bullet')

        doc.add_paragraph()
        p = doc.add_paragraph();
        run = p.add_run("💡 Jak poprawić:");
        run.bold = True;
        run.font.color.rgb = RGBColor(50, 100, 200)
        doc.add_paragraph("1. Rozbuduj treść do minimum 600-800 słów (artykuły: 1500+ słów)", style='List Number')
        doc.add_paragraph("2. Dodaj wartość: praktyczne porady, case studies, przykłady", style='List Number')
        doc.add_paragraph("3. Strukturyzuj: nagłówki H2/H3, listy punktowane, wyróżnienia", style='List Number')
        doc.add_paragraph("4. Multimedialność: obrazy, infografiki, video (zaliczają się do 'treści')",
                          style='List Number')
        doc.add_paragraph("5. FAQ: dodaj sekcję pytań i odpowiedzi (boost dla SEO i użytkownika)", style='List Number')
    else:
        p = doc.add_paragraph();
        run = p.add_run("✅ Brak problemów z thin content");
        run.font.color.rgb = RGBColor(0, 150, 0);
        run.bold = True
        doc.add_paragraph("Wszystkie strony mają odpowiednią ilość treści (>300 słów).")

    doc.add_page_break()
    doc.add_heading('11. 🔒 BEZPIECZEŃSTWO (Security Headers)', 1)

    doc.add_paragraph(
        "Security headers to specjalne nagłówki HTTP, które serwer wysyła do przeglądarki, informując ją 'jak ma się zachować' z punktu widzenia bezpieczeństwa. Przykłady: 'wymuszaj HTTPS', 'nie pozwalaj na osadzenie w iframe', 'blokuj ładowanie zasobów z nieznanych źródeł'. To pierwsza linia obrony przed atakami hackerskimi.")
    doc.add_paragraph()

    p = doc.add_paragraph();
    run = p.add_run('Stan bezpieczeństwa: ');
    run.bold = True
    if summary['avg_security_score'] >= 80:
        sec_status = "✅ Dobry";
        sec_color = RGBColor(0, 150, 0)
    elif summary['avg_security_score'] >= 60:
        sec_status = "⚠️ Średni";
        sec_color = RGBColor(200, 100, 0)
    elif summary['avg_security_score'] >= 40:
        sec_status = "🟠 Słaby";
        sec_color = RGBColor(200, 50, 0)
    else:
        sec_status = "🔴 Krytyczny";
        sec_color = RGBColor(200, 0, 0)

    run = doc.add_paragraph().add_run(f"Status Security: {sec_status} | Śr. ocena: {summary['avg_security_score']}%")
    run.font.color.rgb = sec_color;
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph("Co ryzykujesz bez security headers:")
    doc.add_paragraph("• Ataki XSS (Cross-Site Scripting) - wstrzykiwanie złośliwego kodu", style='List Bullet')
    doc.add_paragraph("• Clickjacking - nakładanie niewidocznych przycisków na Twoją stronę", style='List Bullet')
    doc.add_paragraph("• Man-in-the-middle - przechwytywanie danych użytkowników", style='List Bullet')
    doc.add_paragraph("• Obniżone zaufanie użytkowników (przeglądarki pokazują warningi)", style='List Bullet')
    doc.add_paragraph("• Gorsze pozycjonowanie (Google preferuje bezpieczne strony)", style='List Bullet')
    doc.add_paragraph()

    has_security_issues = (
            summary['pages_poor_security'] > 0 or
            summary['pages_missing_security_headers'] > 0 or
            summary['pages_with_mixed_content'] > 0
    )

    if not has_security_issues:
        p = doc.add_paragraph();
        run = p.add_run("✅ Brak poważnych problemów bezpieczeństwa.");
        run.font.color.rgb = RGBColor(0, 150, 0);
        run.bold = True
    else:
        if issues['poor_security']:
            p = doc.add_paragraph();
            run = p.add_run(f"🟠 Słabe bezpieczeństwo: {len(issues['poor_security'])} stron");
            run.bold = True
            doc.add_paragraph(
                "Strony z oceną security <50% mają krytyczne braki w podstawowych nagłówkach zabezpieczających. Poniżej endpointy z największymi problemami:")
            doc.add_paragraph()

            for item in issues['poor_security'][:10]:
                url = item['url']
                sec_pct = item.get('security_percentage', 0)
                missing = item.get('missing_headers', [])

                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"{url} ")
                run.font.size = Pt(9)
                run2 = p.add_run(f"({sec_pct:.1f}%)\n")
                run2.font.color.rgb = RGBColor(200, 0, 0)
                run2.font.size = Pt(9)

                if missing:
                    run3 = p.add_run(f"  Brakuje: {', '.join(missing[:4])}")
                    run3.font.size = Pt(9)
                    run3.italic = True

            if len(issues['poor_security']) > 10:
                doc.add_paragraph(f"...oraz {len(issues['poor_security']) - 10} innych stron")
            doc.add_paragraph()

        if issues['missing_security_headers']:
            p = doc.add_paragraph();
            run = p.add_run(f"🟡 Braki w security headers: {len(issues['missing_security_headers'])} stron");
            run.bold = True
            doc.add_paragraph("Strony z mniej niż 3 nagłówkami bezpieczeństwa (z 7 możliwych).")
            doc.add_paragraph()

            doc.add_paragraph()
            p = doc.add_paragraph();
            run = p.add_run("Najważniejsze security headers (czego brakuje):");
            run.bold = True;
            run.font.size = Pt(11)

            security_headers_info = [
                ("HSTS (Strict-Transport-Security)",
                 "Wymusza połączenia HTTPS przez określony czas. Bez niego: możliwy atak man-in-the-middle, gdzie hacker przechwytuje dane przesyłane przez HTTP. Google od 2014 preferuje strony HTTPS w rankingu."),

                ("CSP (Content-Security-Policy)",
                 "Określa, z jakich źródeł można ładować zasoby (skrypty, obrazy, CSS). Bez niego: łatwe ataki XSS (Cross-Site Scripting), gdzie hacker wstrzykuje złośliwy kod JavaScript na Twoją stronę."),

                ("X-Frame-Options",
                 "Zapobiega osadzeniu Twojej strony w iframe na innej witrynie. Bez niego: atak clickjacking - hacker nakłada niewidoczny iframe z Twoją stroną na swoją, użytkownik myśli że klika jedno, a klika drugie."),

                ("X-Content-Type-Options",
                 "Blokuje 'MIME sniffing' przeglądarek - zgadywanie typu pliku. Bez niego: przeglądarka może potraktować plik tekstowy jako wykonywalny kod i uruchomić go (atak)."),

                ("Referrer-Policy",
                 "Kontroluje, ile informacji o źródle ruchu jest przekazywane innym stronom. Bez niego: pełny URL (z parametrami, tokenami) może wyciec do zewnętrznych serwisów przez header Referer."),

                ("Permissions-Policy",
                 "Ogranicza dostęp do API przeglądarki (kamera, mikrofon, GPS, etc.). Bez niego: strony osadzone w iframe mogą prosić o dostęp do wrażliwych zasobów użytkownika."),
            ]

            for header_name, description in security_headers_info:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"{header_name}\n")
                run.bold = True
                run.font.size = Pt(10)
                run2 = p.add_run(f"  {description}")
                run2.font.size = Pt(9)
                run2.italic = True

            doc.add_paragraph()
            p = doc.add_paragraph();
            run = p.add_run("Przykładowe endpointy z brakującymi headerami:");
            run.bold = True;
            run.font.size = Pt(11)

            for item in issues['missing_security_headers'][:10]:
                url = item['url']
                headers_count = item['headers_count']
                missing = item.get('missing_critical', [])

                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"{url} ")
                run.font.size = Pt(9)
                run2 = p.add_run(f"({headers_count}/7 headerów)\n")
                run2.font.color.rgb = RGBColor(200, 100, 0)
                run2.font.size = Pt(9)

                if missing:
                    run3 = p.add_run(f"  Brak: {', '.join(missing[:3])}")
                    run3.font.size = Pt(9)
                    run3.italic = True

            if len(issues['missing_security_headers']) > 10:
                doc.add_paragraph(f"...oraz {len(issues['missing_security_headers']) - 10} innych stron")
            doc.add_paragraph()

        if issues['mixed_content']:
            p = doc.add_paragraph();
            run = p.add_run(f"🟡 Mixed content: {len(issues['mixed_content'])} stron");
            run.bold = True
            doc.add_paragraph(
                "Strony HTTPS zawierają zasoby ładowane przez HTTP (obrazki, CSS, JS). Przeglądarka pokazuje ostrzeżenie 'Not Secure' mimo certyfikatu SSL. Użytkownik traci zaufanie, Google obniża ranking.")
            doc.add_paragraph()

            p = doc.add_paragraph();
            run = p.add_run("Strony z mixed content:");
            run.bold = True;
            run.font.size = Pt(11)
            for url in issues['mixed_content'][:10]:
                doc.add_paragraph(f"• {url}", style='List Bullet')
            if len(issues['mixed_content']) > 10:
                doc.add_paragraph(f"...oraz {len(issues['mixed_content']) - 10} innych")

    if SHOW_REMEDIATIONS:
        doc.add_page_break()
        doc.add_heading('🔧 Jak wdrożyć security headers?', 2)
        apache_code = '''Header always set Strict-Transport-Security "max-age=31536000; includeSubDomains"
Header always set X-Frame-Options "DENY"
Header always set X-Content-Type-Options "nosniff"
Header always set Referrer-Policy "strict-origin-when-cross-origin"
Header always set Permissions-Policy "geolocation=(), microphone=(), camera=()"'''
        q = doc.add_paragraph();
        r = q.add_run(apache_code);
        r.font.name = 'Courier New';
        r.font.size = Pt(9)
        doc.add_paragraph()
        nginx_code = '''add_header Strict-Transport-Security "max-age=31536000; includeSubDomains" always;
add_header X-Frame-Options "DENY" always;
add_header X-Content-Type-Options "nosniff" always;
add_header Referrer-Policy "strict-origin-when-cross-origin" always;
add_header Permissions-Policy "geolocation=(), microphone=(), camera=()" always;'''
        q = doc.add_paragraph();
        r = q.add_run(nginx_code);
        r.font.name = 'Courier New';
        r.font.size = Pt(9)
    doc.add_page_break()
    doc.add_heading('12. 📚 Legenda i objaśnienia - Słownik dla biznesu', 1)

    doc.add_paragraph(
        "Poniżej znajdziesz wyjaśnienia wszystkich terminów z raportu - bez żargonu IT, językiem biznesowym. Każdy punkt to konkretny problem lub szansa, która wpływa na Twój ruch, konwersję i sprzedaż.")
    doc.add_paragraph()

    # SEO
    p = doc.add_paragraph();
    run = p.add_run("SEO (Search Engine Optimization)");
    run.bold = True;
    run.font.size = Pt(12);
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "To fundamenty techniczne, które decydują, czy Google w ogóle 'widzi' Twoją stronę i rozumie, o czym jest. SEO to nie magia - to konkretne elementy: czy strona ma tytuł (tag <title>), czy działa na mobile, czy ma poprawne nagłówki H1/H2/H3, czy nie ma błędów 404. Bez tego Google nie wie, co indeksować i gdzie Cię pokazać. Problem: jeśli masz 200 stron bez Title, Google sam wymyśla tytuły - często nietrafiające w sedno, co obniża CTR (Click-Through Rate) o 30-50%. Rezultat: mniej kliknięć z wyszukiwarki = mniej ruchu = mniej leadów/sprzedaży. Inwestycja w SEO to nie koszt, to zwrot: każda złotówka wydana na SEO zwraca się wielokrotnie w postaci darmowego ruchu organicznego.")
    doc.add_paragraph()

    # AEO
    p = doc.add_paragraph();
    run = p.add_run("AEO (Answer Engine Optimization)");
    run.bold = True;
    run.font.size = Pt(12);
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "To ewolucja SEO w erze AI. Użytkownicy coraz częściej pytają Google konkretnymi pytaniami ('jak wybrać laptop do pracy?'), a Google stara się odpowiedzieć BEZPOŚREDNIO - bez konieczności klikania w wyniki. To tzw. featured snippet (pozycja 0) lub rich answer. AEO to formatowanie treści tak, żeby Google mógł z nich łatwo wyciągnąć odpowiedź: listy punktowane, tabele, FAQ, jasne definicje, konkretne liczby. Problem: jeśli Twoja strona ma 'ścianę tekstu' bez struktury, Google pominie Cię i wybierze konkurencję, która ma ładnie sformatowane FAQ. Rezultat: tracisz pozycję 0, która generuje 30-40% kliknięć z pierwszej strony wyników. Rozwiązanie: dodaj sekcje FAQ na każdej ważnej podstronie, używaj list numerowanych (1. 2. 3.) i punktowanych (•), zaznaczaj kluczowe fragmenty.")
    doc.add_paragraph()

    # GEO
    p = doc.add_paragraph();
    run = p.add_run("GEO (Generative Engine Optimization)");
    run.bold = True;
    run.font.size = Pt(12);
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "To najnowszy wymiar - optymalizacja pod AI (ChatGPT, Gemini, Bing Chat, SGE Google). Te systemy nie tylko wyszukują, ale GENERUJĄ odpowiedzi, syntetyzując informacje z wielu źródeł. GEO to sygnały jakości: E-E-A-T (eksperckość, autorytet, zaufanie), dane strukturalne Schema.org, linki do źródeł, daty publikacji, nazwiska autorów. Problem: AI preferuje 'wiarygodne' źródła - jeśli Twoja strona nie ma autora, dat, certyfikatów, AI pominie Cię i zacytuje konkurencję. Rezultat: tracisz ruch z nowego kanału (AI chatboty), który do 2025 będzie generował 20-30% zapytań. Co więcej: Google coraz bardziej opiera się na AI (algorytm RankBrain, BERT, MUM), więc GEO wpływa też na klasyczne pozycjonowanie. Rozwiązanie: dodaj ramki 'O autorze' z bio, linkuj do badań/statystyk, wdróż Schema.org, aktualizuj daty publikacji.")
    doc.add_paragraph()

    # E-E-A-T
    p = doc.add_paragraph();
    run = p.add_run("E-E-A-T (Experience, Expertise, Authoritativeness, Trustworthiness)");
    run.bold = True;
    run.font.size = Pt(12);
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "To nie ranking factor w klasycznym sensie, ale zestaw sygnałów, które Quality Raterzy Google (ludzie oceniający jakość wyników) sprawdzają ręcznie. Algorytm uczy się na ich ocenach. E-E-A-T sprawdza: czy autor jest ekspertem (widoczne imię, nazwisko, bio, doświadczenie), czy treść jest aktualna (data publikacji, ostatnia aktualizacja), czy strona jest autorytetem (linki z .edu/.gov, cytowania w mediach), czy można zaufać (HTTPS, dane kontaktowe, opinie). Problem: jeśli Twój blog nie ma autorów, dat, linków do źródeł - Google traktuje go jako 'low quality' i schodzi w rankingu. To szczególnie ważne dla YMYL (Your Money Your Life) - tematów związanych z finansami, zdrowiem, prawem. Rezultat: spadek pozycji o 10-20 miejsc = utrata 50-70% ruchu. Rozwiązanie jest prosta: dodaj 'O autorze', daty, certyfikaty, linki do badań.")
    doc.add_paragraph()

    # Schema.org
    p = doc.add_paragraph();
    run = p.add_run("Schema.org (Dane strukturalne JSON-LD)");
    run.bold = True;
    run.font.size = Pt(12);
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "To 'język', którym mówisz Google o zawartości strony w sposób ustrukturyzowany. Zamiast 'Google, zgadnij co to za strona', mówisz: 'to jest artykuł, napisany 2024-05-15, autor Jan Kowalski, kategoria Marketing'. Schema to JSON-LD - kod w formacie JSON osadzony w <script type='application/ld+json'>. Problem: bez Schema Google 'zgaduje' co jest czym - i często się myli. Ze Schema możesz mieć rich snippets: gwiazdki ocen (Product), FAQ rozwijane w SERP (FAQPage), breadcrumbs (BreadcrumbList), karuzele wydarzeń (Event). Rezultat: CTR rośnie o 20-30% (kto nie kliknie w 5 gwiazdek?), pozycja 0 (featured snippet) staje się osiągalna, AI (GEO) lepiej rozumie kontekst. Rozwiązanie: wdróż Schema dla strony głównej (Organization/LocalBusiness), produktów (Product), artykułów (Article), FAQ (FAQPage).")
    doc.add_paragraph()

    # NAP (Local SEO)
    p = doc.add_paragraph();
    run = p.add_run("NAP (Name, Address, Phone) - Local SEO");
    run.bold = True;
    run.font.size = Pt(12);
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "To podstawa lokalnego pozycjonowania. NAP to nazwa firmy, adres i telefon - dane, które muszą być SPÓJNE na stronie, Google Maps, Facebooku, wizytówkach. Google sprawdza tę spójność i na jej podstawie decyduje, czy wyświetlić Cię w Local Pack (3 wyniki na mapie powyżej organicznych). Problem: jeśli na stronie masz 'ul. Kwiatowa 5', a na Google Maps 'Kwiatowa 5 lok. 2' - Google nie wie, który jest prawidłowy i obniża Twoją widoczność w wynikach lokalnych. To krytyczne dla biznesów stacjonarnych: restauracje, salony, sklepy, biura. Rezultat: tracisz ruch z zapytań typu 'fryzjer Warszawa Mokotów' (60% zapytań lokalnych kończy się wizytą w lokalu w ciągu 24h). Rozwiązanie: zunifikuj dane NAP wszędzie (identyczne!), dodaj Schema LocalBusiness, osadź mapę Google, stwórz dedykowaną podstronę /kontakt z pełnymi danymi.")
    doc.add_paragraph()

    # Mobile-first indexing
    p = doc.add_paragraph();
    run = p.add_run("Mobile-first indexing (Indeksowanie mobile-first)");
    run.bold = True;
    run.font.size = Pt(12);
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "Od marca 2018 Google NAJPIERW analizuje wersję mobilną strony, dopiero potem desktopową. To nie jest 'dodatek', to główny sposób indeksowania. Jeśli Twoja strona nie ma meta viewport (<meta name='viewport' content='width=device-width, initial-scale=1.0'>), Google traktuje ją jako 'nieresponsywną' i OBNIŻA ranking. Problem: 60-70% ruchu to mobile - jeśli strona źle wygląda na smartfonie, użytkownik wychodzi (bounce rate 80%+), Google to widzi i schodzi Cię w wynikach. To błędne koło: zła mobilność = wysoki bounce = niższy ranking = mniej ruchu. Rezultat: tracisz 50-60% potencjalnego ruchu organicznego. Rozwiązanie: dodaj meta viewport do WSZYSTKICH podstron, przetestuj stronę na urządzeniach mobilnych (Google Mobile-Friendly Test), upewnij się że przyciski są klikalne (min. 48x48px), tekst czytelny (min. 16px).")
    doc.add_paragraph()

    # Security Headers
    p = doc.add_paragraph();
    run = p.add_run("Security Headers (Nagłówki bezpieczeństwa HTTP)");
    run.bold = True;
    run.font.size = Pt(12);
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "To nagłówki HTTP, które serwer wysyła do przeglądarki, informując 'jak się zachować' z punktu widzenia bezpieczeństwa. Przykłady: HSTS (wymuszaj HTTPS), CSP (blokuj ładowanie zasobów z nieznanych źródeł), X-Frame-Options (nie pozwalaj na osadzenie w iframe). Problem: bez nich strona jest podatna na ataki - XSS (wstrzykiwanie złośliwego kodu), clickjacking (nakładanie niewidocznych przycisków), man-in-the-middle (przechwytywanie danych). To nie tylko 'techniczny problem' - jeśli hacker przejmie Twoją stronę i wykradnie dane klientów, możesz mieć RODO-we kary (do 20 mln EUR lub 4% obrotu). Rezultat: oprócz ryzyka prawnego, użytkownicy tracą zaufanie (przeglądarki pokazują 'Not Secure'), Google obniża ranking (preferuje bezpieczne strony od 2014). Rozwiązanie: wdróż 7 kluczowych headerów (HSTS, CSP, X-Frame-Options, X-Content-Type-Options, Referrer-Policy, Permissions-Policy, X-XSS-Protection) - to konfiguracja serwera, 15 minut pracy developera.")
    doc.add_paragraph()

    # Canonical
    p = doc.add_paragraph();
    run = p.add_run("Canonical (Tag kanoniczny)");
    run.bold = True;
    run.font.size = Pt(12);
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "To tag <link rel='canonical' href='https://twoja-domena.pl/strona'>, który mówi Google: 'ta wersja strony jest główna, indeksuj ją'. Problem: często ta sama treść jest dostępna pod wieloma URL-ami - np. /produkt, /produkt?ref=facebook, /produkt?utm_source=newsletter. Bez canonical Google traktuje je jako osobne strony, dzieli 'moc' linków między nimi, wybiera losowo którą pokazać w wynikach. Rezultat: masz 5 wersji tej samej strony, każda na pozycji 50+ zamiast jednej na pozycji 10. To 'kanibalizacja' - strony konkurują ze sobą zamiast się wspierać. Dodatkowo Google może uznać stronę za 'duplicate content' i obniżyć ją w rankingu lub całkowicie wyindeksować. Rozwiązanie: dodaj canonical do KAŻDEJ podstrony wskazujący na 'główną' wersję URL (bez parametrów). Dla produktów z wariantami (kolor, rozmiar) - canonical na podstawową wersję.")
    doc.add_paragraph()

    # Thin content
    p = doc.add_paragraph();
    run = p.add_run("Thin content (Cienka treść)");
    run.bold = True;
    run.font.size = Pt(12);
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "To strony z bardzo małą ilością tekstu - poniżej 300 słów. Google zakłada, że krótka strona = mało wartości = niska jakość. Oczywiście są wyjątki (landing page produktowy z video, strona kontaktu), ale generalnie: im więcej merytorycznej treści, tym lepiej. Problem: algorytm Google Panda (od 2011) karze 'thin content' - strony schodzą w rankingu, tracą indeksację, a w skrajnych przypadkach cała domena może być obniżona (penalty całej witryny). To szczególnie dotyczy sklepów e-commerce z autogenerowanymi opisami produktów (50 słów skopiowanych od producenta). Rezultat: tracisz pozycje, ruch, sprzedaż. Badania pokazują: strony w TOP 10 mają średnio 1500-2000 słów, w TOP 3 - 2500+ słów. Rozwiązanie: rozbuduj treść do min. 600-800 słów (artykuły blogowe: 1500-2500 słów), dodaj wartość (porady, case studies, FAQ), strukturyzuj (nagłówki H2/H3, listy, wyróżnienia).")
    doc.add_paragraph()

    # Open Graph & Twitter Cards
    p = doc.add_paragraph();
    run = p.add_run("Open Graph & Twitter Cards (Meta tagi social media)");
    run.bold = True;
    run.font.size = Pt(12);
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "To meta tagi, które kontrolują, jak Twój link wygląda po udostępnieniu na Facebook, LinkedIn, WhatsApp, X (Twitter). Bez nich platforma sama wybiera obrazek (często logo albo losowe zdjęcie), tytuł (losowy fragment <title>) i opis (pierwsze zdanie tekstu). Problem: tracisz kontrolę nad przekazem marketingowym - zamiast atrakcyjnego obrazka produktu z ceną, Facebook pokazuje małe logo firmy. CTR z social media spada o 50-70%. To ważne, bo social media to potężne źródło ruchu - jeden viralowy post na LinkedIn może przynieść 10 000+ odwiedzin. Rezultat: mniej kliknięć z social = mniej ruchu = mniej konwersji. Dodatkowo: profesjonalny wygląd linku buduje zaufanie do marki. Rozwiązanie: dodaj 4 podstawowe tagi OG (og:title, og:description, og:image, og:url) i 4 tagi Twitter Card (twitter:card, twitter:title, twitter:description, twitter:image). Obrazek: 1200x630px, format JPG/PNG, do 5 MB.")
    doc.add_paragraph()

    doc.add_paragraph()
    p = doc.add_paragraph();
    run = p.add_run("💡 Podsumowanie:");
    run.bold = True;
    run.font.size = Pt(12);
    run.font.color.rgb = RGBColor(50, 100, 200)
    doc.add_paragraph(
        "Każdy z tych elementów to nie 'techniczny detal', ale konkretny wpływ na ruch, konwersję i sprzedaż. SEO to nie koszt, to inwestycja zwracająca się wielokrotnie. Priorytet: zacznij od błędów krytycznych (4xx, brak Title/Description, brak viewport), potem ostrzeżenia (canonical, schema), na końcu optymalizacje (E-E-A-T, content).")

    doc.add_page_break()
    if USE_AI_SUMMARY and OPENAI_API_KEY:
        ai_summary_text = generate_ai_summary(summary, issues)
        if ai_summary_text:
            doc.add_heading('13. 🤖 AI-Powered Executive Summary', 1)
            lines = ai_summary_text.split('\n')
            current_paragraph = None
            for line in lines:
                line = line.strip()
                if not line:
                    current_paragraph = None
                    continue
                if line.startswith('###'):
                    heading_text = line.replace('###', '').strip()
                    doc.add_heading(heading_text, level=3)
                    current_paragraph = None
                elif line.startswith('- ') or line.startswith('* '):
                    list_text = line[2:].strip()
                    doc.add_paragraph(list_text, style='List Bullet')
                    current_paragraph = None
                elif re.match(r'^\d+\.', line):
                    list_text = re.sub(r'^\d+\.\s*', '', line)
                    doc.add_paragraph(list_text, style='List Number')
                    current_paragraph = None
                else:
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph()
                    parts = re.split(r'(\*\*.*?\*\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = current_paragraph.add_run(part[2:-2]);
                            run.bold = True
                        else:
                            current_paragraph.add_run(part)
    doc.add_paragraph('_' * 100)
    p = doc.add_paragraph();
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Raport wygenerowany przez Audytor SEO/AEO/GEO Enhanced Edition\n')
    run.font.size = Pt(9);
    run.italic = True;
    run.font.color.rgb = RGBColor(120, 120, 120)
    p = doc.add_paragraph();
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Data: {summary["generated_at"]}');
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(150, 150, 150)
    doc.save(word_path)
    print(f"✅ Raport Word zapisany: {word_path}")


def save_reports(all_pages: Dict[str, Any], start_url: str, output_dir: str):
    duplicates = find_duplicates(all_pages)
    issues = analyze_issues(all_pages)
    summary = calculate_summary(all_pages, issues, duplicates)
    ai_summary_text = ""
    if USE_AI_SUMMARY and OPENAI_API_KEY:
        print("🤖 Generuję AI Summary...")
        ai_summary_text = generate_ai_summary(summary, issues)
        if ai_summary_text:
            print("✅ AI Summary wygenerowane")
    json_path = os.path.join(output_dir, "raport_szczegolowy.json")
    payload = {"summary": summary, "ai_summary": ai_summary_text, "pages": all_pages, "issues": issues,
               "duplicates": duplicates}
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    print(f"✅ JSON zapisany: {json_path}")
    if pd is not None:
        csv_path = os.path.join(output_dir, "raport_tabela.csv")
        rows = []
        for u, v in all_pages.items():
            rows.append({
                "url": u,
                "excluded": v.get("is_excluded", False),
                "status": v.get("status"),
                "title": v.get("title"),
                "title_length": v.get("meta_scores", {}).get("title_length"),
                "meta_description": v.get("meta_description"),
                "desc_length": v.get("meta_scores", {}).get("desc_length"),
                "canonical": v.get("canonical"),
                "h1_count": v.get("h1_count"),
                "h2_count": v.get("h2_count"),
                "mobile_friendly": v.get("is_mobile_friendly"),
                "has_og_image": v.get("has_og_image"),
                "has_twitter_card": v.get("has_twitter_card"),
                "schema_types": ",".join(v.get("jsonld_types", [])),
                "schema_count": v.get("schema_count"),
                "eeat_score": v.get("eeat_signals", {}).get("eeat_percentage"),
                "nap_score": v.get("nap_signals", {}).get("nap_score"),
                "word_count": v.get("word_count"),
                "img_total": v.get("img_total"),
                "img_without_alt": v.get("img_without_alt"),
                "error": v.get("error"),
                "has_ssl": v.get("security", {}).get("has_ssl"),
                "security_score": v.get("security", {}).get("security_percentage"),
                "security_level": v.get("security", {}).get("security_level"),
                "security_headers_count": v.get("security", {}).get("headers_count"),
                "has_hsts": v.get("security", {}).get("security_checks", {}).get("hsts", {}).get("present"),
                "has_csp": v.get("security", {}).get("security_checks", {}).get("content_security_policy", {}).get(
                    "present"),
                "has_mixed_content": v.get("security", {}).get("has_mixed_content"),
            })
        df = pd.DataFrame(rows)
        df.to_csv(csv_path, index=False, encoding="utf-8")
        print(f"✅ CSV zapisany: {csv_path}")
    word_path = os.path.join(output_dir, "raport_dla_klienta.docx")
    create_word_report(all_pages, summary, issues, duplicates, word_path)


if __name__ == "__main__":
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    OUTPUT_DIR = f"audyt_{timestamp}"
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print("=" * 80)
    print("🚀 AUDYTOR SEO/AEO/GEO - ENHANCED EDITION".center(80))
    print("=" * 80)
    print(f"\n📁 Katalog wyników: {OUTPUT_DIR}")
    print(f"🌐 Audytowana strona: {START_URL}")
    print(f"⚙️  MAX_PAGES={MAX_PAGES}, MAX_DEPTH={MAX_DEPTH}, CONCURRENCY={CONCURRENCY}")
    print(f"🚫 Wykluczono m.in. /cdn-cgi/*")
    if USE_PAGESPEED:
        print(f"📊 PageSpeed Insights: WŁĄCZONY")
    else:
        print(f"📊 PageSpeed Insights: WYŁĄCZONY")
    if USE_AI_SUMMARY:
        if OPENAI_API_KEY:
            print(f"🤖 AI Summary: WŁĄCZONY (model: {OPENAI_MODEL})")
        else:
            print(f"🤖 AI Summary: WYŁĄCZONY (brak OPENAI_API_KEY)")
    else:
        print(f"🤖 AI Summary: WYŁĄCZONY")
    print("\n" + "=" * 80)
    print()
    t0 = time.time()
    data = asyncio.run(crawl(START_URL))
    save_reports(data, START_URL, OUTPUT_DIR)
    print("\n" + "=" * 80)
    print("📊 AUDYT ZAKOŃCZONY".center(80))
    print("=" * 80)
    print(f"\n⏱️  Czas wykonania: {time.time() - t0:.1f}s")
    print(f"\n📁 Wyniki: {OUTPUT_DIR}/")
    print("   • raport_dla_klienta.docx")
    print("   • raport_szczegolowy.json")
    print("   • raport_tabela.csv")
    print("\n✅ Gotowe!")