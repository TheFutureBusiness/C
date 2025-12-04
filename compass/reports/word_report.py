"""
Generator raportu Word (DOCX)

Obsługuje wielojęzyczność - język raportu konfigurowany w config.REPORT_LANGUAGE
"""
from typing import Dict, Any
import re

from compass.config import SHOW_REMEDIATIONS, USE_AI_SUMMARY, OPENAI_API_KEY, REPORT_LANGUAGE
from compass.reports.translations import t

try:
    from docx import Document as WordDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False
    print("⚠️  Brak biblioteki python-docx. Raport Word nie będzie wygenerowany.")


def add_hyperlink(paragraph, text, url):
    """
    Dodaje hyperlink do paragrafu w dokumencie Word.

    Args:
        paragraph: Paragraf w dokumencie Word
        text: Tekst linku
        url: URL linku
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_section_heading(doc, text, level=1, icon=None, color=RGBColor(31, 71, 136)):
    """
    Dodaje spójny nagłówek sekcji z opcjonalną ikonką emoji.
    """
    if icon:
        full_text = f"{icon} {text}"
    else:
        full_text = text
    heading = doc.add_heading(full_text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color
    return heading


def add_status_line(doc, label, status_text, color, extra=None):
    """
    Dodaje wyróżnioną linię statusu (np. Status SEO, Status Security).
    """
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r2 = p.add_run(status_text)
    r2.bold = True
    r2.font.color.rgb = color
    if extra:
        p.add_run(f" – {extra}")


def create_word_report(all_pages: Dict[str, Any], summary: Dict[str, Any],
                       issues: Dict[str, Any], duplicates: Dict, word_path: str):
    if not HAS_DOCX:
        print("⚠️  Pomijam generowanie raportu Word")
        return

    # Język raportu
    lang = REPORT_LANGUAGE

    doc = WordDocument()

    # =========================
    # OKŁADKA RAPORTU + KPI
    # =========================
    title = doc.add_heading(t('report_title', lang), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(31, 71, 136)
    title_run.font.size = Pt(34)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{summary['start_url']}\n")
    run.font.size = Pt(18)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{t('audit_date', lang)}: {summary['generated_at']}")
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(120, 120, 120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(t('recipients', lang))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph()

    # MAŁA TABELKA KPI NA OKŁADCE
    kpi_table = doc.add_table(rows=2, cols=3)
    kpi_table.style = 'Light Shading Accent 1'

    hdr = kpi_table.rows[0].cells
    hdr[0].text = "🎯 Hauptergebnis"
    hdr[1].text = "✅ Technische Stabilität"
    hdr[2].text = "📱 Mobilität"

    kpis = [
        ("Audit-Ergebnis", f"{summary['overall_score']}/100", summary['overall_grade']),
        ("Seiten OK (200)", summary['pages_ok'], f"von {summary['pages_analyzed']} analysierten"),
        ("Mobile-friendly", f"{summary['mobile_percentage']}%", f"{summary['mobile_friendly_pages']} Seiten"),
    ]

    row = kpi_table.rows[1].cells
    for cell, (label, value, sub) in zip(row, kpis):
        p = cell.paragraphs[0]
        r1 = p.add_run(f"{label}\n")
        r1.bold = True
        r2 = p.add_run(f"{value}\n")
        r2.font.size = Pt(14)
        r2.bold = True
        r3 = p.add_run(str(sub))
        r3.font.size = Pt(9)
        r3.italic = True

    doc.add_paragraph()

    # WYNIK AUDYTU jako wyróżniony box
    box = doc.add_paragraph()
    run = box.add_run(f"AUDIT-ERGEBNIS: {summary['overall_score']}/100  ({summary['overall_grade']})")
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 71, 136)

    doc.add_paragraph()

    # SKŁAD OCENY OGÓLNEJ
    add_section_heading(doc, 'Zusammensetzung der Gesamtbewertung', level=3)
    scoring_table = doc.add_table(rows=1, cols=3)
    scoring_table.style = 'Light Grid Accent 1'
    hdr = scoring_table.rows[0].cells
    hdr[0].text = 'Kategorie'
    hdr[1].text = 'Wert (x von y)'
    hdr[2].text = 'Beschreibung'

    # Obliczenie wartości dla każdej kategorii
    pages = max(1, summary['pages_analyzed'])

    # Dostępność
    availability_pct = (summary['pages_ok'] / pages) * 100
    availability_value = round(availability_pct * 0.30)
    availability_text = f"{availability_value} z 30 ({availability_pct:.1f}%)"

    # Meta tagi
    meta_ok_pages = pages - (summary['missing_title'] + summary['missing_description'])
    meta_pct = max(0.0, (meta_ok_pages - 0.25 * (summary['title_issues'] + summary['description_issues'])) / pages * 100)
    meta_value = round(meta_pct * 0.15)
    meta_text = f"{meta_value} z 15 ({meta_pct:.1f}%)"

    # Mobile-friendly
    mobile_value = round(summary['mobile_percentage'] * 0.15)
    mobile_text = f"{mobile_value} z 15 ({summary['mobile_percentage']:.1f}%)"

    # Schema.org
    schema_pct = (summary['pages_with_schema'] / pages) * 100
    schema_value = round(schema_pct * 0.10)
    schema_text = f"{schema_value} z 10 ({schema_pct:.1f}%)"

    # E-E-A-T
    eeat_value = round(summary['avg_eeat_score'] * 0.10)
    eeat_text = f"{eeat_value} z 10 ({summary['avg_eeat_score']:.1f}%)"

    # Bezpieczeństwo
    security_pct = summary['avg_security_score']
    if summary['pages_no_ssl'] > 0:
        security_pct = max(0.0, security_pct - 10)
    security_value = round(security_pct * 0.20)
    security_text = f"{security_value} z 20 ({security_pct:.1f}%)"

    scoring_components = [
        ('Verfügbarkeit (HTTP 200)', availability_text, 'Prozentsatz der Seiten ohne HTTP-Fehler (4xx, 5xx)'),
        ('Meta-Tags', meta_text, 'Vorhandensein und Qualität von Title und Description'),
        ('Mobile-friendly', mobile_text, 'Responsive Design und Meta Viewport'),
        ('Schema.org', schema_text, 'Strukturierte Daten JSON-LD'),
        ('E-E-A-T', eeat_text, 'Expertise, Autorität, Vertrauen'),
        ('Sicherheit', security_text, 'HTTPS, Security Headers, kein Mixed Content'),
    ]

    for category, weight, description in scoring_components:
        row = scoring_table.add_row().cells
        row[0].text = category
        row[1].text = weight
        row[2].text = description

    doc.add_paragraph()

    # Najważniejsze problemy z licznikami X/Y i opisami
    add_section_heading(doc, 'Wichtigste Probleme', level=2, icon='⚠️')

    # Słownik z opisami problemów
    problem_descriptions = {
        "pages_with_errors": "Seiten mit HTTP-Fehlercodes (4xx, 5xx). Sie verhindern die Indexierung und führen zu Traffic-Verlust.",
        "missing_title": "Jede Seite benötigt ein eindeutiges <title>-Tag (50–60 Zeichen) für bessere Sichtbarkeit in Suchergebnissen.",
        "missing_description": "Meta Description (150–160 Zeichen) ist der erste Kontakt des Nutzers mit Ihrer Seite in Google-Ergebnissen.",
        "missing_canonical": "Das Canonical-Tag verhindert Probleme mit doppelten Inhalten und hilft Google, die richtige Seitenversion zu wählen.",
        "pages_without_viewport": "Fehlendes Meta Viewport - erfordert manuelle Überprüfung der Darstellung auf Mobilgeräten.",
        "pages_without_schema": "Strukturierte Daten (Schema.org) helfen Google, Inhalte besser zu verstehen und Rich Snippets anzuzeigen.",
        "pages_poor_security": "Schwache Sicherheit (<50%). Fehlende Security Headers gefährden Nutzer und reduzieren das Vertrauen.",
        "thin_content_pages": "Seiten mit weniger als 300 Wörtern. Google bevorzugt wertvolle, detaillierte Inhalte.",
        "pages_weak_eeat": "Schwache E-E-A-T-Signale (<50%). Fügen Sie Autor, Veröffentlichungsdatum, Zertifikate und Links zu vertrauenswürdigen Quellen hinzu.",
    }

    top_issues = [
        ("HTTP-Fehler (4xx/5xx)", summary["pages_with_errors"], summary["pages_analyzed"], "pages_with_errors"),
        ("Fehlender Title", summary["missing_title"], summary["pages_analyzed"], "missing_title"),
        ("Fehlende Meta Description", summary["missing_description"], summary["pages_analyzed"], "missing_description"),
        ("Fehlende Canonical", summary["missing_canonical"], summary["pages_analyzed"], "missing_canonical"),
        ("Fehlendes Meta Viewport (Mobile)", summary["pages_without_viewport"], summary["pages_analyzed"],
         "pages_without_viewport"),
        ("Fehlendes Schema.org", summary["pages_without_schema"], summary["pages_analyzed"], "pages_without_schema"),
        ("Schwache Sicherheit (<50%)", summary["pages_poor_security"], summary["pages_analyzed"],
         "pages_poor_security"),
        ("Dünner Inhalt (<300 Wörter)", summary["thin_content_pages"], summary["pages_analyzed"], "thin_content_pages"),
        ("Schwaches E-E-A-T (<50%)", summary["pages_weak_eeat"], summary["pages_analyzed"], "pages_weak_eeat"),
    ]

    for label, problem_count, total_pages, key in sorted(top_issues, key=lambda kv: kv[1], reverse=True)[:6]:
        if problem_count > 0:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"{label}: {problem_count}/{total_pages}\n")
            run.bold = True
            if key in problem_descriptions:
                run2 = p.add_run(problem_descriptions[key])
                run2.font.size = Pt(10)
                run2.italic = True
                run2.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_page_break()

    # =========================
    # SPIS TREŚCI
    # =========================
    add_section_heading(doc, 'Inhaltsverzeichnis', 1, icon='📋')
    toc_items = [
        "1. Executive Summary – Kernzahlen",
        "2. Prioritäten (von kritisch bis gering)",
        "3. Meta-Tag-Analyse",
        "4. Technische SEO-Analyse",
        "5. Mobilität und Responsive Design",
        "6. Open Graph und Twitter Cards",
        "7. Strukturierte Daten (Schema.org)",
        "8. E-E-A-T",
        "9. Local SEO (NAP)",
        "10. Inhaltsqualität",
        "11. Sicherheit (Security Headers)",
        "12. Legende und Erläuterungen",
        "13. KI-gestützte Zusammenfassung",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
    doc.add_page_break()

    # =========================
    # 1. EXECUTIVE SUMMARY
    # =========================
    add_section_heading(doc, '1. Executive Summary – Kernzahlen', 1, icon='📊')
    add_section_heading(doc, 'Schnelle SEO-Zusammenfassung', 2, icon='📌')

    p = doc.add_paragraph()
    run = p.add_run('Status der Website auf einen Blick:\n')
    run.font.size = Pt(11)
    run.bold = True

    # Problemy KRYTYCZNE: 4xx, brak SSL, brak Title, brak Description
    seo_critical = summary['pages_with_errors'] + summary['pages_no_ssl'] + summary['missing_title'] + summary[
        'missing_description']

    # Ostrzeżenia: problemy z długością, canonical, schema, itp.
    seo_warnings = summary['title_issues'] + summary['description_issues'] + summary['missing_canonical'] + summary[
        'pages_without_schema']

    if seo_critical == 0 and seo_warnings < 10:
        seo_status = "✅ Ausgezeichnet!"
        seo_color = RGBColor(0, 150, 0)
        seo_text = "Kleinere Optimierungen – SEO-Grundlagen in gutem Zustand."
    elif seo_critical < 5 and seo_warnings < 30:
        seo_status = "⚠️ Verbesserungsbedarf"
        seo_color = RGBColor(200, 100, 0)
        seo_text = f"Kritische Probleme: {seo_critical} | Warnungen: {seo_warnings}"
    else:
        seo_status = "🔴 Erfordert Aufmerksamkeit!"
        seo_color = RGBColor(200, 0, 0)
        seo_text = (
            f"Kritische Probleme: {seo_critical} (4xx-Fehler, fehlendes SSL/Title/Description) | "
            f"Warnungen: {seo_warnings} (Meta-Länge, Canonical, Schema)"
        )

    add_status_line(doc, "SEO-Status", seo_status, seo_color, extra=seo_text)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('🎯 Top 3 Prioritäten für den nächsten Sprint')
    run.font.size = Pt(12)
    run.bold = True

    priorities = []
    total_pages = summary['pages_analyzed']

    if summary['pages_with_errors'] > 0:
        pct = round((summary['pages_with_errors'] / total_pages) * 100, 1)
        priorities.append({
            'text': f"Beheben Sie {summary['pages_with_errors']}/{total_pages} Seiten mit HTTP 4xx-Fehlern ({pct}%)",
            'desc': "Seiten nicht erreichbar für Nutzer und Google-Bots",
            'count': summary['pages_with_errors']
        })

    if summary['missing_title'] > 0:
        pct = round((summary['missing_title'] / total_pages) * 100, 1)
        priorities.append({
            'text': f"Fügen Sie Title zu {summary['missing_title']}/{total_pages} Seiten hinzu ({pct}%)",
            'desc': "Fehlender Titel = Unsichtbarkeit in Google-Ergebnissen",
            'count': summary['missing_title']
        })

    if summary['missing_description'] > 0 and len(priorities) < 3:
        pct = round((summary['missing_description'] / total_pages) * 100, 1)
        priorities.append({
            'text': f"Fügen Sie Meta Description zu {summary['missing_description']}/{total_pages} Seiten hinzu ({pct}%)",
            'desc': "Beeinflusst die CTR (Click-Through Rate) aus der Suche",
            'count': summary['missing_description']
        })

    if summary['pages_without_schema'] > 0 and len(priorities) < 3:
        pct = round((summary['pages_without_schema'] / total_pages) * 100, 1)
        priorities.append({
            'text': f"Fügen Sie Schema.org zu {summary['pages_without_schema']}/{total_pages} Seiten hinzu ({pct}%)",
            'desc': "Keine Rich Snippets in Google (Sterne, FAQ, Breadcrumbs)",
            'count': summary['pages_without_schema']
        })

    if summary['missing_canonical'] > 0 and len(priorities) < 3:
        pct = round((summary['missing_canonical'] / total_pages) * 100, 1)
        priorities.append({
            'text': f"Fügen Sie Canonical zu {summary['missing_canonical']}/{total_pages} Seiten hinzu ({pct}%)",
            'desc': "Verhindert Probleme mit doppelten Inhalten",
            'count': summary['missing_canonical']
        })

    if summary['pages_without_viewport'] > 0 and len(priorities) < 3:
        pct = round((summary['pages_without_viewport'] / total_pages) * 100, 1)
        priorities.append({
            'text': f"Überprüfen Sie mobile Darstellung von {summary['pages_without_viewport']}/{total_pages} Seiten ohne Viewport ({pct}%)",
            'desc': "Erfordert manuelle Überprüfung auf Mobilgeräten",
            'count': summary['pages_without_viewport']
        })

    # Sortujemy po count (malejąco) i bierzemy top 3
    priorities.sort(key=lambda x: x['count'], reverse=True)

    for i, pr in enumerate(priorities[:3], 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f"{i}. {pr['text']}\n")
        run.bold = True
        run2 = p.add_run(f"   Geschäftlicher Nutzen: {pr['desc']}")
        run2.font.size = Pt(10)
        run2.italic = True
        run2.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    # Krótkie "co zyskasz"
    p = doc.add_paragraph()
    run = p.add_run("Was Sie durch Umsetzung der Empfehlungen gewinnen:")
    run.bold = True
    doc.add_paragraph("• mehr organischer Traffic von Google", style='List Bullet')
    doc.add_paragraph("• höhere CTR aus Suchergebnissen und Social Media", style='List Bullet')
    doc.add_paragraph("• bessere Sicherheit und Nutzervertrauen", style='List Bullet')

    doc.add_paragraph()
    stats_table = doc.add_table(rows=1, cols=3)
    stats_table.style = 'Light Grid Accent 1'
    hdr = stats_table.rows[0].cells
    hdr[0].text = 'Metrik'
    hdr[1].text = 'Wert'
    hdr[2].text = 'Beschreibung'

    # Funkcja pomocnicza do dodawania wierszy z opisem
    def add_metric_row(metric_name, value, description="", emoji=None):
        row = stats_table.add_row().cells
        label = f"{emoji} {metric_name}" if emoji else metric_name
        row[0].text = label
        row[1].text = str(value)
        row[2].text = description

    # Metryki z opisami
    add_metric_row(
        'Gescannte Seiten',
        f"{summary['pages_crawled']}",
        "Gesamtzahl der während des Crawlings entdeckten URLs",
        emoji="🌐"
    )

    add_metric_row(
        'Analysierte Seiten',
        f"{summary['pages_analyzed']} ({summary['pages_excluded']} ausgeschlossen)",
        "Seiten die einer SEO-Analyse unterzogen wurden (ohne technische URLs wie /cdn-cgi/*)",
        emoji="📄"
    )

    add_metric_row(
        'Seiten OK (200)',
        f"✅ {summary['pages_ok']}",
        "Korrekt funktionierende Seiten, die HTTP-Code 200 zurückgeben",
        emoji="✅"
    )

    add_metric_row(
        'Seiten mit Fehlern (4xx)',
        f"🔴 {summary['pages_with_errors']}",
        "Client-Fehler (404 Not Found, 403 Forbidden etc.) – blockieren Indexierung",
        emoji="🚫"
    )

    add_metric_row(
        'Fehlender Title',
        f"{summary['missing_title']}",
        "Seiten ohne <title>-Tag – entscheidend für Ranking und CTR",
        emoji="🧾"
    )

    add_metric_row(
        'Fehlende Description',
        f"{summary['missing_description']}",
        "Seiten ohne <meta name=\"description\"> – beeinflusst Snippet in Google",
        emoji="📝"
    )

    add_metric_row(
        'Title-Probleme (Länge)',
        f"{summary['title_issues']}",
        "Title zu kurz (<30 Zeichen) oder zu lang (>65) – Google kann abschneiden",
        emoji="📏"
    )

    add_metric_row(
        'Description-Probleme (Länge)',
        f"{summary['description_issues']}",
        "Description außerhalb des Bereichs 120–165 Zeichen – kann abgeschnitten oder zu kurz sein",
        emoji="📐"
    )

    add_metric_row(
        'Title-Duplikate',
        f"{summary['duplicate_titles']}",
        "Verschiedene Seiten mit identischem Titel – verwirrt Google und Nutzer",
        emoji="🔁"
    )

    add_metric_row(
        'Description-Duplikate',
        f"{summary['duplicate_descriptions']}",
        "Verschiedene Seiten mit derselben Meta Description – verringert Einzigartigkeit",
        emoji="🔁"
    )

    add_metric_row(
        'Fehlende Canonical',
        f"{summary['missing_canonical']}",
        "Fehlendes <link rel=\"canonical\">-Tag – führt zu Problemen mit doppelten Inhalten",
        emoji="🏷️"
    )

    add_metric_row(
        'Fehlende H1',
        f"{summary['missing_h1']}",
        "Seite ohne Hauptüberschrift <h1> – Schlüsselelement für Struktur und SEO",
        emoji="🔤"
    )

    add_metric_row(
        'Mehrere H1',
        f"{summary['multiple_h1']}",
        "Seite mit >1 H1-Überschrift – kann Suchmaschinen verwirren",
        emoji="⚠️"
    )

    add_metric_row(
        'Bilder ohne ALT (Gesamtzahl)',
        f"{summary['total_images_without_alt']}",
        "Bilder ohne alt-Attribut – Problem für SEO, Barrierefreiheit und Google Bilder",
        emoji="🖼️"
    )

    add_metric_row(
        'Mobile-friendly',
        f"{summary['mobile_friendly_pages']} ({summary['mobile_percentage']}%)",
        "Für Mobilgeräte optimierte Seiten (Responsive Design)",
        emoji="📱"
    )

    add_metric_row(
        'Fehlendes Meta Viewport',
        f"{summary['pages_without_viewport']}",
        "Seiten ohne <meta name=\"viewport\"> – skalieren nicht auf Mobile",
        emoji="🔍"
    )

    add_metric_row(
        'Fehlendes Open Graph',
        f"{summary['pages_without_og']}",
        "Seiten ohne OG Meta-Tags (Facebook, LinkedIn) – hässliche Link-Vorschau",
        emoji="📢"
    )

    add_metric_row(
        'Fehlende Twitter Cards',
        f"{summary['pages_without_twitter']}",
        "Seiten ohne Twitter Card – keine attraktive Vorschau auf X (Twitter)",
        emoji="🐦"
    )

    add_metric_row(
        'Seiten mit strukturierten Daten',
        f"{summary['pages_with_schema']}",
        "Anzahl der Seiten mit Schema.org (JSON-LD) – ermöglicht Rich Snippets in Google",
        emoji="🔗"
    )

    add_metric_row(
        'Fehlendes Schema',
        f"{summary['pages_without_schema']}",
        "Seiten ohne Schema.org – Sie verlieren Sterne, FAQ, Breadcrumbs in Ergebnissen",
        emoji="⚠️"
    )

    add_metric_row(
        'Durchschn. Schema-Typen/Seite',
        f"{summary['avg_schema_types']}",
        "Wie viele verschiedene Schema-Typen auf der Seite sind (Article, Product, FAQ, etc.)",
        emoji="📚"
    )

    add_metric_row(
        'Durchschn. E-E-A-T',
        f"{summary['avg_eeat_score']}%",
        "Experience, Expertise, Authoritativeness, Trustworthiness – Google-Qualitätssignale",
        emoji="🏆"
    )

    add_metric_row(
        'Schwaches E-E-A-T',
        f"{summary['pages_weak_eeat']}",
        "Seiten mit niedrigem E-E-A-T (<50%): fehlender Autor, Datum, Zertifikate, Quellen",
        emoji="⚠️"
    )

    add_metric_row(
        'Local NAP OK',
        f"{summary['local_optimized_pages']}",
        "Seiten mit korrekten NAP-Daten (Name, Address, Phone) – wichtig für lokale Unternehmen",
        emoji="📍"
    )

    add_metric_row(
        'Schwaches Local SEO',
        f"{summary['pages_poor_local_seo']}",
        "Seiten ohne NAP, Schema LocalBusiness, Kartenlinks – schlecht für lokales SEO",
        emoji="📉"
    )

    add_metric_row(
        'Dünner Inhalt (<300 Wörter)',
        f"{summary['thin_content_pages']}",
        "Seiten mit sehr kurzem Inhalt – Google kann als Low-Quality einstufen",
        emoji="✂️"
    )

    add_metric_row(
        'Durchschn. Sicherheit',
        f"{summary['avg_security_score']}%",
        "Durchschnittliches Sicherheitsniveau (HTTPS + Security Headers). {:.1f}% ist sehr niedrig".format(
            summary['avg_security_score']),
        emoji="🔒"
    )

    add_metric_row(
        'Schwache Sicherheit',
        f"{summary['pages_poor_security']}",
        "Seiten mit Bewertung <50%: fehlende kritische Header (HSTS, CSP, X-Frame-Options)",
        emoji="🛑"
    )

    add_metric_row(
        'Fehlende Security Headers',
        f"{summary['pages_missing_security_headers']}",
        "Seiten mit <3 Security Headers. Prüfen Sie: HSTS, CSP, X-Frame-Options, X-Content-Type-Options, Referrer-Policy",
        emoji="🧱"
    )

    add_metric_row(
        'Mixed Content',
        f"{summary['pages_with_mixed_content']}",
        "HTTPS-Seiten mit HTTP-Ressourcen (Bilder, Skripte) – Warnung im Browser",
        emoji="⚡"
    )

    add_metric_row(
        'AUDIT-ERGEBNIS',
        f"{summary['overall_score']}/100 ({summary['overall_grade']})",
        "Gesamtbewertung: Verfügbarkeit (30%) + Meta (15%) + Mobile (15%) + Schema (10%) + E-E-A-T (10%) + Sicherheit (20%)",
        emoji="🏁"
    )

    doc.add_page_break()

    # =========================
    # 2. PRIORYTETY
    # =========================
    add_section_heading(doc, '2. Prioritäten (von kritisch bis gering)', 1, icon='🎯')

    if issues['critical_errors']:
        add_section_heading(doc, 'HTTP-Fehler (4xx)', 2, icon='🔴')
        p = doc.add_paragraph()
        total = summary['pages_analyzed']
        count = len(issues['critical_errors'])
        pct = round((count / total) * 100, 1) if total > 0 else 0
        p.add_run(f"Gefunden: {count}/{total} Seiten mit 4xx-Fehlern ({pct}%).").bold = True
        doc.add_paragraph(
            "Diese Seiten sind für Nutzer und Google-Bots nicht erreichbar, was zu Folgendem führt:\n"
            "• Verlust von organischem Traffic\n"
            "• Negativer Einfluss auf UX\n"
            "• Probleme bei der Indexierung"
        )
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Problematische Adressen:")
        run.bold = True
        for err in issues['critical_errors'][:20]:
            status_code = err.get('status', 'N/A')
            error_msg = err.get('error', '')
            if error_msg and error_msg != '':
                doc.add_paragraph(f"• {err['url']} – Status: {status_code} ({error_msg})", style='List Bullet')
            else:
                doc.add_paragraph(f"• {err['url']} – Status: {status_code}", style='List Bullet')
        if len(issues['critical_errors']) > 20:
            doc.add_paragraph(f"...sowie {len(issues['critical_errors']) - 20} weitere")

    if issues['missing_title'] or issues['title_issues'] or issues['missing_description'] or issues['description_issues'] or issues['missing_canonical']:
        add_section_heading(doc, 'Meta & Canonical', 2, icon='🟠')

        # ===== META TITLE =====
        if issues['missing_title']:
            p = doc.add_paragraph()
            total = summary['pages_analyzed']
            count = len(issues['missing_title'])
            pct = round((count / total) * 100, 1) if total > 0 else 0
            run = p.add_run(f"Fehlende Meta Title: {count}/{total} Seiten ({pct}%)")
            run.bold = True
            doc.add_paragraph(
                "Das <title>-Tag ist das erste Element, das Nutzer in Google-Ergebnissen sehen. Sein Fehlen bedeutet:"
            )
            doc.add_paragraph("• Keine Kontrolle darüber, was Google in den SERPs anzeigt", style='List Bullet')
            doc.add_paragraph("• Niedrigere CTR (Click-Through Rate)", style='List Bullet')
            doc.add_paragraph("• Schlechteres Ranking", style='List Bullet')
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("Seiten ohne Title:")
            run.bold = True
            for url in issues['missing_title'][:15]:
                doc.add_paragraph(f"• {url}", style='List Bullet')
            if len(issues['missing_title']) > 15:
                doc.add_paragraph(f"...sowie {len(issues['missing_title']) - 15} weitere")
            doc.add_paragraph()

        if issues['title_issues']:
            # Rozdzielenie na za krótkie i za długie
            too_short = [item for item in issues['title_issues'] if item.get('too_short')]
            too_long = [item for item in issues['title_issues'] if item.get('too_long')]

            if too_short:
                p = doc.add_paragraph()
                total = summary['pages_analyzed']
                count = len(too_short)
                pct = round((count / total) * 100, 1) if total > 0 else 0
                run = p.add_run(f"Meta Title zu kurz (<30 Zeichen): {count}/{total} Seiten ({pct}%)")
                run.bold = True
                doc.add_paragraph(
                    "Ein zu kurzer Titel nutzt den verfügbaren Platz in Google-Ergebnissen (50-60 Zeichen) nicht voll aus. "
                    "Sie verlieren die Möglichkeit, Schlüsselwörter einzubauen und die Aufmerksamkeit der Nutzer zu erregen."
                )
                doc.add_paragraph()
                p = doc.add_paragraph()
                run = p.add_run("Seiten mit zu kurzem Title:")
                run.bold = True
                for item in too_short[:15]:
                    url = item['url']
                    title = item.get('title', '')[:80]
                    length = item.get('length', 0)
                    doc.add_paragraph(f"• {url}\n  Title ({length} Zeichen): \"{title}\"", style='List Bullet')
                if len(too_short) > 15:
                    doc.add_paragraph(f"...sowie {len(too_short) - 15} weitere")
                doc.add_paragraph()

            if too_long:
                p = doc.add_paragraph()
                total = summary['pages_analyzed']
                count = len(too_long)
                pct = round((count / total) * 100, 1) if total > 0 else 0
                run = p.add_run(f"Meta Title zu lang (>65 Zeichen): {count}/{total} Seiten ({pct}%)")
                run.bold = True
                doc.add_paragraph(
                    "Ein zu langer Titel wird in Google-Ergebnissen abgeschnitten (angezeigt werden ~50-60 Zeichen). "
                    "Wichtige Informationen sind möglicherweise für Nutzer nicht sichtbar."
                )
                doc.add_paragraph()
                p = doc.add_paragraph()
                run = p.add_run("Seiten mit zu langem Title:")
                run.bold = True
                for item in too_long[:15]:
                    url = item['url']
                    title = item.get('title', '')[:100]
                    length = item.get('length', 0)
                    doc.add_paragraph(f"• {url}\n  Title ({length} Zeichen): \"{title}...\"", style='List Bullet')
                if len(too_long) > 15:
                    doc.add_paragraph(f"...sowie {len(too_long) - 15} weitere")
                doc.add_paragraph()

        # ===== META DESCRIPTION =====
        if issues['missing_description']:
            p = doc.add_paragraph()
            total = summary['pages_analyzed']
            count = len(issues['missing_description'])
            pct = round((count / total) * 100, 1) if total > 0 else 0
            run = p.add_run(f"Fehlende Meta Description: {count}/{total} Seiten ({pct}%)")
            run.bold = True
            doc.add_paragraph("Meta Description ist die 'Einladung' zum Klicken in Suchergebnissen. Ohne sie:")
            doc.add_paragraph("• Google generiert selbst eine Beschreibung (oft nicht treffend)", style='List Bullet')
            doc.add_paragraph("• Sie verlieren die Kontrolle über die Marketing-Botschaft", style='List Bullet')
            doc.add_paragraph("• CTR kann um 30–40% sinken", style='List Bullet')
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("Seiten ohne Meta Description:")
            run.bold = True
            for url in issues['missing_description'][:15]:
                doc.add_paragraph(f"• {url}", style='List Bullet')
            if len(issues['missing_description']) > 15:
                doc.add_paragraph(f"...sowie {len(issues['missing_description']) - 15} weitere")
            doc.add_paragraph()

        if issues['description_issues']:
            # Rozdzielenie na za krótkie i za długie
            too_short_desc = [item for item in issues['description_issues'] if item.get('too_short')]
            too_long_desc = [item for item in issues['description_issues'] if item.get('too_long')]

            if too_short_desc:
                p = doc.add_paragraph()
                total = summary['pages_analyzed']
                count = len(too_short_desc)
                pct = round((count / total) * 100, 1) if total > 0 else 0
                run = p.add_run(f"Meta Description zu kurz (<120 Zeichen): {count}/{total} Seiten ({pct}%)")
                run.bold = True
                doc.add_paragraph(
                    "Eine zu kurze Beschreibung nutzt den verfügbaren Platz (120-165 Zeichen) nicht aus und überzeugt Nutzer möglicherweise nicht zum Klicken."
                )
                doc.add_paragraph()
                p = doc.add_paragraph()
                run = p.add_run("Seiten mit zu kurzer Description:")
                run.bold = True
                for item in too_short_desc[:15]:
                    url = item['url']
                    desc = item.get('description', '')
                    length = item.get('length', 0)
                    doc.add_paragraph(f"• {url}\n  Description ({length} Zeichen): \"{desc}\"", style='List Bullet')
                if len(too_short_desc) > 15:
                    doc.add_paragraph(f"...sowie {len(too_short_desc) - 15} weitere")
                doc.add_paragraph()

            if too_long_desc:
                p = doc.add_paragraph()
                total = summary['pages_analyzed']
                count = len(too_long_desc)
                pct = round((count / total) * 100, 1) if total > 0 else 0
                run = p.add_run(f"Meta Description zu lang (>165 Zeichen): {count}/{total} Seiten ({pct}%)")
                run.bold = True
                doc.add_paragraph(
                    "Eine zu lange Beschreibung wird in Google-Ergebnissen abgeschnitten. Wichtige Informationen können verborgen bleiben."
                )
                doc.add_paragraph()
                p = doc.add_paragraph()
                run = p.add_run("Seiten mit zu langer Description:")
                run.bold = True
                for item in too_long_desc[:15]:
                    url = item['url']
                    desc = item.get('description', '')
                    length = item.get('length', 0)
                    doc.add_paragraph(f"• {url}\n  Description ({length} Zeichen): \"{desc}...\"", style='List Bullet')
                if len(too_long_desc) > 15:
                    doc.add_paragraph(f"...sowie {len(too_long_desc) - 15} weitere")
                doc.add_paragraph()

        # ===== CANONICAL =====
        if issues['missing_canonical']:
            p = doc.add_paragraph()
            total = summary['pages_analyzed']
            count = len(issues['missing_canonical'])
            pct = round((count / total) * 100, 1) if total > 0 else 0
            run = p.add_run(f"Fehlende Canonical: {count}/{total} Seiten ({pct}%)")
            run.bold = True
            doc.add_paragraph(
                "Das Canonical-Tag zeigt Google, welche Seitenversion die 'Hauptversion' ist. Sein Fehlen führt zu:"
            )
            doc.add_paragraph(
                "• Problemen mit doppelten Inhalten (Google weiß nicht, welche Version indexiert werden soll)",
                style='List Bullet'
            )
            doc.add_paragraph("• Streuung der Link-'Power' zwischen Duplikaten", style='List Bullet')
            doc.add_paragraph("• Schwächerem Ranking aller Versionen", style='List Bullet')
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("Seiten ohne Canonical:")
            run.bold = True
            for url in issues['missing_canonical'][:15]:
                doc.add_paragraph(f"• {url}", style='List Bullet')
            if len(issues['missing_canonical']) > 15:
                doc.add_paragraph(f"...sowie {len(issues['missing_canonical']) - 15} weitere")
            doc.add_paragraph()

    if issues['poor_security'] or issues['missing_security_headers']:
        add_section_heading(doc, 'Sicherheit – niedriges Niveau / fehlende Header', 2, icon='🟠')

        if issues['poor_security']:
            p = doc.add_paragraph()
            total = summary['pages_analyzed']
            count = len(issues['poor_security'])
            pct = round((count / total) * 100, 1) if total > 0 else 0
            run = p.add_run(f"Schwache Sicherheit: {count}/{total} Seiten ({pct}%)")
            run.bold = True
            doc.add_paragraph(
                "Seiten mit Sicherheitsbewertung <50% haben Mängel bei grundlegenden Sicherheitsheadern:"
            )
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("Beispiele für Seiten mit schwacher Sicherheit:")
            run.bold = True
            for item in issues['poor_security'][:10]:
                url = item['url']
                sec_pct = item.get('security_percentage', 0)
                missing = item.get('missing_headers', [])
                if missing:
                    missing_str = ", ".join(missing[:3])
                    doc.add_paragraph(f"• {url} ({sec_pct}%) – Fehlend: {missing_str}", style='List Bullet')
                else:
                    doc.add_paragraph(f"• {url} ({sec_pct}%)", style='List Bullet')
            if len(issues['poor_security']) > 10:
                doc.add_paragraph(f"...sowie {len(issues['poor_security']) - 10} weitere")
            doc.add_paragraph()

        if issues['missing_security_headers']:
            p = doc.add_paragraph()
            total = summary['pages_analyzed']
            count = len(issues['missing_security_headers'])
            pct = round((count / total) * 100, 1) if total > 0 else 0
            run = p.add_run(
                f"Fehlende Security Headers: {count}/{total} Seiten ({pct}%)"
            )
            run.bold = True
            doc.add_paragraph(
                "Seiten mit weniger als 3 Sicherheitsheadern sind anfällig für Angriffe. Fehlende Header bedeuten:"
            )
            doc.add_paragraph("• Leichtere Durchführung von XSS- und Clickjacking-Angriffen", style='List Bullet')
            doc.add_paragraph("• Keine Erzwingung von HTTPS (Man-in-the-Middle möglich)", style='List Bullet')
            doc.add_paragraph("• Verringertes Vertrauen von Nutzern und Google", style='List Bullet')
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("Wichtige fehlende Header:")
            run.bold = True
            p = doc.add_paragraph()
            run = p.add_run("• HSTS: ")
            run.bold = True
            p.add_run("Erzwingt HTTPS, schützt vor Man-in-the-Middle-Angriffen")
            p = doc.add_paragraph()
            run = p.add_run("• CSP: ")
            run.bold = True
            p.add_run("Verhindert XSS-Angriffe (Einschleusen von bösartigem Code)")
            p = doc.add_paragraph()
            run = p.add_run("• X-Frame-Options: ")
            run.bold = True
            p.add_run("Schützt vor Clickjacking (Einbettung der Seite in iframe)")
            doc.add_paragraph()

    if issues['no_viewport'] or issues['no_og_tags'] or issues['no_twitter_cards'] or issues['missing_schema']:
        add_section_heading(doc, 'Mobile / Social / Schema', 2, icon='🟡')

        if issues['no_viewport']:
            p = doc.add_paragraph()
            total = summary['pages_analyzed']
            count = len(issues['no_viewport'])
            pct = round((count / total) * 100, 1) if total > 0 else 0
            run = p.add_run(f"Fehlendes Meta Viewport: {count}/{total} Seiten ({pct}%)")
            run.bold = True
            doc.add_paragraph("Seit 2018 verwendet Google Mobile-First-Indexing. Fehlendes Meta Viewport bedeutet:")
            doc.add_paragraph("• Seite skaliert nicht korrekt auf Smartphones", style='List Bullet')
            doc.add_paragraph("• Google kann Ranking senken (Mobile-First!)", style='List Bullet')
            doc.add_paragraph("• Mobile Nutzer sehen 'Desktop'-Version (schlechte UX)", style='List Bullet')
            doc.add_paragraph()

        if issues['no_og_tags']:
            p = doc.add_paragraph()
            total = summary['pages_analyzed']
            count = len(issues['no_og_tags'])
            pct = round((count / total) * 100, 1) if total > 0 else 0
            run = p.add_run(f"Fehlende Open Graph Tags: {count}/{total} Seiten ({pct}%)")
            run.bold = True
            doc.add_paragraph(
                "Open Graph sind Meta-Tags, die von Facebook, LinkedIn, WhatsApp zur Generierung von Link-Vorschauen verwendet werden."
            )
            doc.add_paragraph()

        if issues['no_twitter_cards']:
            p = doc.add_paragraph()
            total = summary['pages_analyzed']
            count = len(issues['no_twitter_cards'])
            pct = round((count / total) * 100, 1) if total > 0 else 0
            run = p.add_run(f"Fehlende Twitter Cards: {count}/{total} Seiten ({pct}%)")
            run.bold = True
            doc.add_paragraph("Twitter Cards sind das OG-Äquivalent für die Plattform X (ehemals Twitter).")
            doc.add_paragraph()

        if issues['missing_schema']:
            p = doc.add_paragraph()
            total = summary['pages_analyzed']
            count = len(issues['missing_schema'])
            pct = round((count / total) * 100, 1) if total > 0 else 0
            run = p.add_run(f"Fehlendes Schema.org: {count}/{total} Seiten ({pct}%)")
            run.bold = True
            doc.add_paragraph(
                "Schema.org (JSON-LD) ist die 'Sprache', mit der Sie Google über den Inhalt Ihrer Seite informieren. Ohne sie:"
            )
            doc.add_paragraph("• Verlieren Sie Rich Snippets (Sterne, FAQ, Breadcrumbs)", style='List Bullet')
            doc.add_paragraph("• Schwieriger, Featured Snippet (Position 0) zu erreichen", style='List Bullet')
            doc.add_paragraph("• Google versteht Inhaltskontext schlechter", style='List Bullet')
            doc.add_paragraph()

    if issues['weak_eeat'] or issues['thin_content']:
        add_section_heading(doc, 'E-E-A-T & Inhalt', 2, icon='🟡')

        if issues['weak_eeat']:
            p = doc.add_paragraph()
            total = summary['pages_analyzed']
            count = len(issues['weak_eeat'])
            pct = round((count / total) * 100, 1) if total > 0 else 0
            run = p.add_run(f"Schwaches E-E-A-T: {count}/{total} Seiten ({pct}%)")
            run.bold = True
            doc.add_paragraph(
                "E-E-A-T (Experience, Expertise, Authoritativeness, Trustworthiness) ist ein Satz von Qualitätssignalen für Google."
            )
            doc.add_paragraph()

        if issues['thin_content']:
            p = doc.add_paragraph()
            total = summary['pages_analyzed']
            count = len(issues['thin_content'])
            pct = round((count / total) * 100, 1) if total > 0 else 0
            run = p.add_run(f"Dünner Inhalt: {count}/{total} Seiten ({pct}%)")
            run.bold = True
            doc.add_paragraph(
                "Seiten mit weniger als 300 Wörtern können von Google als 'dünn' (Low-Quality) eingestuft werden."
            )
            doc.add_paragraph()

    doc.add_page_break()

    # =========================
    # 5. MOBILNOŚĆ
    # =========================
    add_section_heading(doc, '5. Mobilität und Responsive Design', 1, icon='📱')
    p = doc.add_paragraph()
    run = p.add_run('Mobilitätsstatus: ')
    run.bold = True
    if summary['mobile_percentage'] >= 90:
        run = p.add_run(f"✅ {summary['mobile_percentage']}% der Seiten sind mobile-friendly")
        run.font.color.rgb = RGBColor(0, 150, 0)
    elif summary['mobile_percentage'] >= 70:
        run = p.add_run(f"⚠️ {summary['mobile_percentage']}% der Seiten sind mobile-friendly")
        run.font.color.rgb = RGBColor(200, 100, 0)
    else:
        run = p.add_run(f"🔴 {summary['mobile_percentage']}% der Seiten sind mobile-friendly")
        run.font.color.rgb = RGBColor(200, 0, 0)

    doc.add_paragraph()
    doc.add_paragraph(
        "Google analysiert zuerst die mobile Version der Seite (Mobile-First-Indexing). Fehlende Responsiveness bedeutet:"
    )
    doc.add_paragraph("• Rangabfall in mobilen Ergebnissen (60%+ des Traffics ist mobil)", style='List Bullet')
    doc.add_paragraph("• Schlechtere Nutzererfahrung = höhere Bounce Rate", style='List Bullet')
    doc.add_paragraph("• Verlust potenzieller mobiler Kunden", style='List Bullet')

    if issues['no_viewport']:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(f"Seiten ohne Meta Viewport ({len(issues['no_viewport'])}):")
        run.bold = True
        doc.add_paragraph(
            "⚠️ HINWEIS: Das Fehlen von Meta Viewport garantiert nicht, dass die Seite auf Mobilgeräten falsch angezeigt wird. "
            "Dies ist nur ein technischer Indikator. Wir empfehlen eine manuelle Überprüfung der Darstellung auf echten Mobilgeräten "
            "oder mit Tools wie Google PageSpeed Insights, Lighthouse oder Chrome DevTools."
        )
        doc.add_paragraph()

        # Filtrujemy pliki multimedialne (mp4, mp3, jpg, png, pdf, etc.)
        multimedia_extensions = ['.mp4', '.mp3', '.avi', '.mov', '.wmv', '.flv', '.webm', '.mkv',
                                 '.jpg', '.jpeg', '.png', '.gif', '.webp', '.svg', '.pdf', '.zip',
                                 '.rar', '.doc', '.docx', '.xls', '.xlsx']

        filtered_urls = []
        for url in issues['no_viewport']:
            url_lower = url.lower()
            is_multimedia = any(url_lower.endswith(ext) for ext in multimedia_extensions)
            if not is_multimedia:
                filtered_urls.append(url)

        if filtered_urls:
            p = doc.add_paragraph()
            run = p.add_run("HTML-Seiten die Verbesserung benötigen:")
            run.bold = True
            run.font.size = Pt(11)
            for url in filtered_urls[:15]:
                doc.add_paragraph(f"• {url}", style='List Bullet')
            if len(filtered_urls) > 15:
                doc.add_paragraph(f"...sowie {len(filtered_urls) - 15} weitere Seiten")

        # Jeśli są pliki multimedialne, informujemy o nich osobno
        multimedia_count = len(issues['no_viewport']) - len(filtered_urls)
        if multimedia_count > 0:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(f"ℹ️  Zusätzlich {multimedia_count} Multimediadateien ohne Viewport gefunden ")
            run.font.size = Pt(9)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            run2 = p.add_run("(Video/Bilddateien – normal, keine Verbesserung nötig)")
            run2.font.size = Pt(9)
            run2.italic = True
            run2.font.color.rgb = RGBColor(100, 100, 100)

        if SHOW_REMEDIATIONS:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run('So fügen Sie Meta Viewport hinzu:\n')
            run.bold = True
            code = '<meta name="viewport" content="width=device-width, initial-scale=1.0">'
            q = doc.add_paragraph()
            r = q.add_run(code)
            r.font.name = 'Courier New'
            r.font.size = Pt(10)

    doc.add_page_break()

    # =========================
    # 6. OPEN GRAPH / TWITTER
    # =========================
    add_section_heading(doc, '6. Open Graph und Twitter Cards', 1, icon='📢')

    doc.add_paragraph(
        "Social Media ist eine wichtige Traffic-Quelle. Wenn jemand einen Link zu Ihrer Seite auf Facebook, LinkedIn oder X (Twitter) teilt, generieren diese Plattformen eine 'Vorschau' – ein Miniaturbild mit Bild, Titel und Beschreibung. "
        "Dieser erste Eindruck entscheidet, ob der Nutzer klickt."
    )
    doc.add_paragraph()

    if issues['no_og_tags']:
        p = doc.add_paragraph()
        total = summary['pages_analyzed']
        count = len(issues['no_og_tags'])
        pct = round((count / total) * 100, 1) if total > 0 else 0
        run = p.add_run(f"Fehlende Open Graph Tags: {count}/{total} Seiten ({pct}%)")
        run.bold = True
        doc.add_paragraph("Open Graph sind Meta-Tags, die von Facebook, LinkedIn, WhatsApp, Messenger verwendet werden. Ohne sie:")
        doc.add_paragraph("• Plattformen wählen selbst ein Bild (oft unpassend)", style='List Bullet')
        doc.add_paragraph("• Beschreibung kann ein zufälliger Textabschnitt sein", style='List Bullet')
        doc.add_paragraph("• Sie verlieren die Kontrolle über die Marketing-Botschaft", style='List Bullet')
        doc.add_paragraph("• Niedrigere CTR aus Social Media (bis zu 50%!)", style='List Bullet')
        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("Seiten, die Open Graph benötigen:")
        run.bold = True
        run.font.size = Pt(11)
        for item in issues['no_og_tags'][:15]:
            url = item['url']
            has_image = item.get('has_og_image', False)
            has_title = item.get('has_og_title', False)
            has_desc = item.get('has_og_description', False)

            missing_parts = []
            if not has_image:
                missing_parts.append("Bild")
            if not has_title:
                missing_parts.append("Titel")
            if not has_desc:
                missing_parts.append("Beschreibung")

            if missing_parts:
                doc.add_paragraph(f"• {url}\n  Fehlt: {', '.join(missing_parts)}", style='List Bullet')
            else:
                doc.add_paragraph(f"• {url}", style='List Bullet')

        if len(issues['no_og_tags']) > 15:
            doc.add_paragraph(f"...und {len(issues['no_og_tags']) - 15} weitere Seiten")
        doc.add_paragraph()

    if issues['no_twitter_cards']:
        p = doc.add_paragraph()
        total = summary['pages_analyzed']
        count = len(issues['no_twitter_cards'])
        pct = round((count / total) * 100, 1) if total > 0 else 0
        run = p.add_run(f"Fehlende Twitter Cards: {count}/{total} Seiten ({pct}%)")
        run.bold = True
        doc.add_paragraph(
            "Twitter Cards sind das Äquivalent zu Open Graph für die Plattform X (früher Twitter). Sie funktionieren analog – "
            "sie steuern, wie Ihr Link beim Teilen aussieht."
        )
        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("Seiten, die Twitter Cards benötigen:")
        run.bold = True
        run.font.size = Pt(11)
        for url in issues['no_twitter_cards'][:15]:
            doc.add_paragraph(f"• {url}", style='List Bullet')
        if len(issues['no_twitter_cards']) > 15:
            doc.add_paragraph(f"...und {len(issues['no_twitter_cards']) - 15} weitere Seiten")
        doc.add_paragraph()

    if SHOW_REMEDIATIONS:
        example_code = '''<!-- Open Graph -->
<meta property="og:title" content="Tytuł" />
<meta property="og:description" content="Opis" />
<meta property="og:image" content="https://twoja-domena.pl/obraz.jpg" />
<meta property="og:url" content="https://twoja-domena.pl/strona" />
<meta property="og:type" content="website" />
<!-- Twitter Cards -->
<meta name="twitter:card" content="summary_large_image" />
<meta name="twitter:title" content="Tytuł" />
<meta name="twitter:description" content="Opis" />
<meta name="twitter:image" content="https://twoja-domena.pl/obraz.jpg" />'''
        p = doc.add_paragraph()
        run = p.add_run('Przykładowe tagi:')
        run.bold = True
        q = doc.add_paragraph()
        r = q.add_run(example_code)
        r.font.name = 'Courier New'
        r.font.size = Pt(9)

    doc.add_page_break()

    # =========================
    # 7. SCHEMA
    # =========================
    add_section_heading(doc, '7. Strukturierte Daten (Schema.org)', 1, icon='🔗')

    doc.add_paragraph(
        "Schema.org (JSON-LD und Microdata) ist die 'technische Sprache', mit der Sie Google über den Inhalt Ihrer Seite informieren. "
        "Dadurch kann Google Ihre Seite attraktiver in den Suchergebnissen anzeigen – sogenannte Rich Snippets."
    )
    doc.add_paragraph(
        "HINWEIS: Dieser Audit erkennt sowohl JSON-LD als auch Microdata-Formate."
    )
    doc.add_paragraph()

    schema_percentage = (summary['pages_with_schema'] / max(1, summary['pages_analyzed'])) * 100

    if schema_percentage >= 70:
        p = doc.add_paragraph()
        run = p.add_run(f"✅ {summary['pages_with_schema']} Seiten haben strukturierte Daten ({schema_percentage:.1f}%)")
        run.font.color.rgb = RGBColor(0, 150, 0)
        run.bold = True
    else:
        p = doc.add_paragraph()
        run = p.add_run(
            f"⚠️ Nur {summary['pages_with_schema']} Seiten haben strukturierte Daten ({schema_percentage:.1f}%)"
        )
        run.font.color.rgb = RGBColor(200, 100, 0)
        run.bold = True

    doc.add_paragraph()
    doc.add_paragraph("Was Sie ohne Schema.org verlieren:")
    doc.add_paragraph("• Rich Snippets: Sternebewertungen, Produktpreise, aufklappbare FAQs in SERP", style='List Bullet')
    doc.add_paragraph("• Breadcrumbs (Navigationspfad) in Google-Ergebnissen", style='List Bullet')
    doc.add_paragraph("• Featured Snippet (Position 0) – schwieriger ohne Struktur", style='List Bullet')
    doc.add_paragraph("• Produkt-/Artikelkarussells in mobilen Ergebnissen", style='List Bullet')
    doc.add_paragraph("• Besseres Kontextverständnis durch Google (wichtig für AI)", style='List Bullet')
    doc.add_paragraph()

    if issues['missing_schema']:
        p = doc.add_paragraph()
        total = summary['pages_analyzed']
        count = len(issues['missing_schema'])
        pct = round((count / total) * 100, 1) if total > 0 else 0
        run = p.add_run(f"Seiten ohne Schema.org: {count}/{total} ({pct}%)")
        run.bold = True
        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("Beispiele für Seiten, die Schema benötigen:")
        run.bold = True
        run.font.size = Pt(11)
        for url in issues['missing_schema'][:20]:
            doc.add_paragraph(f"• {url}", style='List Bullet')
        if len(issues['missing_schema']) > 20:
            doc.add_paragraph(f"...und {len(issues['missing_schema']) - 20} weitere Seiten")

        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("💡 Empfehlung:")
        run.bold = True
        run.font.color.rgb = RGBColor(50, 100, 200)
        doc.add_paragraph("Prioritär Schema hinzufügen zu:")
        doc.add_paragraph("Startseite (Organization/LocalBusiness)", style='List Bullet')
        doc.add_paragraph("Produktseiten (Product mit Preis und Verfügbarkeit)", style='List Bullet')
        doc.add_paragraph("Blog-Artikel (Article/BlogPosting)", style='List Bullet')
        doc.add_paragraph("FAQ-Seiten (FAQPage)", style='List Bullet')
        doc.add_paragraph("Kundenbewertungen (Review/AggregateRating)", style='List Bullet')

    doc.add_page_break()

    # =========================
    # 8. E-E-A-T
    # =========================
    add_section_heading(doc, '8. E-E-A-T', 1, icon='🏆')

    doc.add_paragraph(
        "E-E-A-T steht für Experience, Expertise, Authoritativeness, Trustworthiness – auf Deutsch: "
        "Erfahrung, Expertise, Autorität, Vertrauenswürdigkeit."
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "WICHTIG: E-E-A-T wird kontextabhängig bewertet. Nicht jede Seite braucht alle Signale:"
    )
    doc.add_paragraph("• Blog/Artikel: Autor-Box und Veröffentlichungsdatum wichtig", style='List Bullet')
    doc.add_paragraph("• Service-Seiten: Autor nicht erforderlich, Kontaktdaten und Expertise wichtiger", style='List Bullet')
    doc.add_paragraph("• Über uns: Zertifikate, Auszeichnungen, Erfahrung besonders relevant", style='List Bullet')
    doc.add_paragraph("• Alle Seiten: HTTPS und Kontaktmöglichkeiten im Footer", style='List Bullet')
    doc.add_paragraph()

    doc.add_paragraph("Was Google bei E-E-A-T bewertet:")
    doc.add_paragraph("• Ist der Autor ein Experte? (sichtbarer Name, Foto, Bio – nur bei Blog)", style='List Bullet')
    doc.add_paragraph("• Ist der Inhalt aktuell? (Veröffentlichungsdatum, Aktualisierungsdatum – nur bei Blog)", style='List Bullet')
    doc.add_paragraph("• Ist die Seite eine Autorität? (Links zu .edu/.gov Quellen)", style='List Bullet')
    doc.add_paragraph("• Kann man vertrauen? (HTTPS, Kontaktdaten, Datenschutz)", style='List Bullet')
    doc.add_paragraph("• Gibt es Bewertungen/Rezensionen? (Social Proof)", style='List Bullet')
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('Durchschnittliche E-E-A-T Bewertung: ')
    run.bold = True
    if summary['avg_eeat_score'] >= 70:
        run = p.add_run(f"✅ {summary['avg_eeat_score']}%")
        run.font.color.rgb = RGBColor(0, 150, 0)
    elif summary['avg_eeat_score'] >= 50:
        run = p.add_run(f"⚠️ {summary['avg_eeat_score']}%")
        run.font.color.rgb = RGBColor(200, 100, 0)
    else:
        run = p.add_run(f"🔴 {summary['avg_eeat_score']}%")
        run.font.color.rgb = RGBColor(200, 0, 0)

    doc.add_paragraph()

    if issues['weak_eeat']:
        p = doc.add_paragraph()
        total = summary['pages_analyzed']
        count = len(issues['weak_eeat'])
        pct = round((count / total) * 100, 1) if total > 0 else 0
        run = p.add_run(f"Seiten mit schwachem E-E-A-T (<50%): {count}/{total} ({pct}%)")
        run.bold = True
        doc.add_paragraph(
            "Diese Seiten haben aus Google-Sicht einen niedrigen Vertrauenswert. Das bedeutet nicht, dass sie 'schlecht' sind – "
            "es fehlen nur einige Qualitätssignale (je nach Seitentyp unterschiedlich gewichtet)."
        )
        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("Beispiele für Seiten mit schwachem E-E-A-T:")
        run.bold = True
        run.font.size = Pt(11)
        for item in issues['weak_eeat'][:15]:
            url = item['url']
            eeat_pct = item.get('eeat_percentage', 0)
            missing = item.get('missing', [])

            missing_readable = []
            for key in missing:
                if key == 'has_author':
                    missing_readable.append("Autor")
                elif key == 'has_date':
                    missing_readable.append("Veröffentlichungsdatum")
                elif key == 'has_expertise_signals':
                    missing_readable.append("Expertise-Signale")
                elif key == 'has_quality_external_links':
                    missing_readable.append("Quellenlinks")
                elif key == 'has_contact_info':
                    missing_readable.append("Kontaktdaten")
                elif key == 'has_reviews':
                    missing_readable.append("Bewertungen")

            if missing_readable:
                doc.add_paragraph(
                    f"• {url} ({eeat_pct:.1f}%)\n  Fehlt: {', '.join(missing_readable[:3])}",
                    style='List Bullet'
                )
            else:
                doc.add_paragraph(f"• {url} ({eeat_pct:.1f}%)", style='List Bullet')

        if len(issues['weak_eeat']) > 15:
            doc.add_paragraph(f"...und {len(issues['weak_eeat']) - 15} weitere Seiten")

        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("💡 E-E-A-T verbessern (kontextabhängig):")
        run.bold = True
        run.font.color.rgb = RGBColor(50, 100, 200)
        doc.add_paragraph("Bei Blog-Artikeln: Autor-Box mit Name, Foto, Bio hinzufügen", style='List Bullet')
        doc.add_paragraph("Bei Blog-Artikeln: Veröffentlichungsdatum und 'Letzte Aktualisierung' anzeigen", style='List Bullet')
        doc.add_paragraph("Überall: Links zu vertrauenswürdigen Quellen (.edu, .gov, Studien)", style='List Bullet')
        doc.add_paragraph("Auf 'Über uns': Zertifikate, Auszeichnungen, Unternehmenserfahrung", style='List Bullet')
        doc.add_paragraph("Im Footer: Sichtbare Kontaktdaten (Telefon, E-Mail, Adresse)", style='List Bullet')

    doc.add_page_break()

    # =========================
    # 9. LOCAL SEO (NAP)
    # =========================
    add_section_heading(doc, '9. Local SEO (NAP)', 1, icon='📍')

    doc.add_paragraph(
        "NAP steht für Name, Address, Phone – also Firmenname, Adresse und Telefon. Das ist die Grundlage für lokales SEO."
    )
    doc.add_paragraph(
        "HINWEIS: NAP-Daten sollten einmal im Footer oder auf der Kontaktseite vorhanden sein. "
        "Wenn sie dort korrekt eingebunden sind, gilt die Website als NAP-optimiert."
    )
    doc.add_paragraph()
    doc.add_paragraph("Google prüft, ob NAP-Daten:",)
    doc.add_paragraph("• Konsistent sind (überall gleich: Website, Google Maps, Facebook, Visitenkarten)", style='List Bullet')
    doc.add_paragraph("• Sichtbar sind (leicht auf der Seite zu finden, z.B. im Footer)", style='List Bullet')
    doc.add_paragraph("• Strukturiert sind (Schema.org Organization/LocalBusiness)", style='List Bullet')
    doc.add_paragraph()

    local_percentage = (summary['local_optimized_pages'] / max(1, summary['pages_analyzed'])) * 100

    if local_percentage >= 50:
        p = doc.add_paragraph()
        run = p.add_run(
            f"✅ {summary['local_optimized_pages']} Seiten NAP-optimiert ({local_percentage:.1f}%)"
        )
        run.font.color.rgb = RGBColor(0, 150, 0)
        run.bold = True
    else:
        p = doc.add_paragraph()
        run = p.add_run(
            f"⚠️ Nur {summary['local_optimized_pages']} Seiten haben NAP ({local_percentage:.1f}%)"
        )
        run.font.color.rgb = RGBColor(200, 100, 0)
        run.bold = True

    doc.add_paragraph()
    doc.add_paragraph("Warum NAP wichtig ist:")
    doc.add_paragraph("• Google Local Pack (3 Ergebnisse auf der Karte) erfordert konsistente Daten", style='List Bullet')
    doc.add_paragraph("• Nutzer, die nach 'Firma + Stadt' suchen, landen auf lokalen Ergebnissen", style='List Bullet')
    doc.add_paragraph("• Vertrauen: sichtbare Telefonnummer und Adresse = höhere Conversion", style='List Bullet')
    doc.add_paragraph("• Voice Search ('Hey Google, finde X in der Nähe') bevorzugt NAP", style='List Bullet')
    doc.add_paragraph()

    if issues['poor_local_seo']:
        p = doc.add_paragraph()
        total = summary['pages_analyzed']
        count = len(issues['poor_local_seo'])
        pct = round((count / total) * 100, 1) if total > 0 else 0
        run = p.add_run(f"Seiten ohne korrektes NAP: {count}/{total} ({pct}%)")
        run.bold = True
        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("Beispiele für Seiten, die Local SEO Verbesserung benötigen:")
        run.bold = True
        run.font.size = Pt(11)
        for item in issues['poor_local_seo'][:15]:
            url = item['url']
            nap_score = item.get('nap_score', 0)
            phones = item.get('phone_numbers', 0)
            has_address = item.get('has_address', False)
            has_schema = item.get('has_local_schema', False)

            issues_list = []
            if phones == 0:
                issues_list.append("keine Telefonnummer")
            if not has_address:
                issues_list.append("keine Adresse")
            if not has_schema:
                issues_list.append("kein Schema LocalBusiness")

            if issues_list:
                doc.add_paragraph(
                    f"• {url} (NAP: {nap_score}/3)\n  Problem: {', '.join(issues_list)}",
                    style='List Bullet'
                )
            else:
                doc.add_paragraph(f"• {url} (NAP: {nap_score}/3)", style='List Bullet')

        if len(issues['poor_local_seo']) > 15:
            doc.add_paragraph(f"...und {len(issues['poor_local_seo']) - 15} weitere Seiten")

        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("💡 Local SEO verbessern:")
        run.bold = True
        run.font.color.rgb = RGBColor(50, 100, 200)
        doc.add_paragraph("Sichtbaren Footer mit: Firmenname, Adresse, Telefon, E-Mail hinzufügen", style='List Bullet')
        doc.add_paragraph("Schema.org Organization/LocalBusiness (JSON-LD) mit vollständigen NAP-Daten implementieren", style='List Bullet')
        doc.add_paragraph("Konsistenz prüfen: gleiche Daten auf Google Maps, Facebook, Website", style='List Bullet')
        doc.add_paragraph("Google Maps (Embed) mit Standortangabe hinzufügen", style='List Bullet')
        doc.add_paragraph("Dedizierte 'Kontakt'-Seite mit vollständigen Daten erstellen", style='List Bullet')

    doc.add_page_break()

    # =========================
    # 10. INHALTSQUALITÄT
    # =========================
    add_section_heading(doc, '10. Inhaltsqualität', 1, icon='📝')

    doc.add_paragraph(
        "Google mag keine 'dünnen' Seiten – also solche mit sehr wenig Text (unter 300 Wörtern). "
        "Je mehr fachlicher Inhalt, desto besser."
    )
    doc.add_paragraph()

    doc.add_paragraph("Risiken bei Thin Content:")
    doc.add_paragraph("• Google kann die Seite als minderwertig einstufen und im Ranking herabsetzen", style='List Bullet')
    doc.add_paragraph("• Schwieriger, Featured Snippet (Position 0) zu erreichen – mehr Kontext erforderlich", style='List Bullet')
    doc.add_paragraph("• Nutzer verlassen die Seite schnell (hohe Absprungrate)", style='List Bullet')
    doc.add_paragraph("• Geringere Chance auf Backlinks (niemand verlinkt einen 100-Wörter-Text)", style='List Bullet')
    doc.add_paragraph()

    if issues['thin_content']:
        p = doc.add_paragraph()
        total = summary['pages_analyzed']
        count = len(issues['thin_content'])
        pct = round((count / total) * 100, 1) if total > 0 else 0
        run = p.add_run(f"Seiten mit Thin Content (<300 Wörter): {count}/{total} ({pct}%)")
        run.bold = True
        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("Seiten, die Inhaltserweiterung benötigen:")
        run.bold = True
        run.font.size = Pt(11)
        for item in issues['thin_content']:
            url = item['url']
            word_count = item.get('word_count', 0)
            text_len = item.get('text_len', 0)

            doc.add_paragraph(f"• {url}\n  Wörter: {word_count}, Zeichen: {text_len}", style='List Bullet')

        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("💡 Verbesserungsvorschläge:")
        run.bold = True
        run.font.color.rgb = RGBColor(50, 100, 200)
        doc.add_paragraph("Inhalt auf mindestens 600–800 Wörter erweitern (Artikel: 1500+ Wörter)", style='List Bullet')
        doc.add_paragraph("Mehrwert bieten: praktische Tipps, Case Studies, Beispiele", style='List Bullet')
        doc.add_paragraph("Strukturieren: H2/H3-Überschriften, Aufzählungen, Hervorhebungen", style='List Bullet')
        doc.add_paragraph("Multimedialität: Bilder, Infografiken, Videos (zählen als 'Inhalt')", style='List Bullet')
        doc.add_paragraph("FAQ: Fragen-und-Antworten-Bereich hinzufügen (SEO-Boost)", style='List Bullet')
    else:
        p = doc.add_paragraph()
        run = p.add_run("✅ Keine Thin Content Probleme")
        run.font.color.rgb = RGBColor(0, 150, 0)
        run.bold = True
        doc.add_paragraph("Alle Seiten haben ausreichend Inhalt (>300 Wörter).")

    doc.add_page_break()

    # =========================
    # 11. BEZPIECZEŃSTWO
    # =========================
    add_section_heading(doc, '11. Sicherheit (Security Headers)', 1, icon='🔒')

    doc.add_paragraph(
        "Security Headers sind spezielle HTTP-Header, die der Server an den Browser sendet und ihm mitteilen, "
        "'wie er sich verhalten soll' im Hinblick auf die Sicherheit."
    )
    doc.add_paragraph()

    # Bestimmen des Sicherheitsstatus basierend auf dem Durchschnitt
    if summary['avg_security_score'] >= 80:
        sec_status = "✅ Ausgezeichnet"
        sec_color = RGBColor(0, 150, 0)
        sec_msg = "Die Website erfüllt hohe Sicherheitsstandards."
    elif summary['avg_security_score'] >= 60:
        sec_status = "🟡 Gut"
        sec_color = RGBColor(200, 180, 0)
        sec_msg = "Grundlegende Sicherheitsmaßnahmen sind vorhanden, aber Verbesserungen möglich."
    elif summary['avg_security_score'] >= 40:
        sec_status = "🟠 Verbesserungswürdig"
        sec_color = RGBColor(200, 100, 0)
        sec_msg = "Einige Security Headers fehlen. Empfohlen: Konfiguration verbessern."
    else:
        sec_status = "🔴 Kritisch"
        sec_color = RGBColor(200, 0, 0)
        sec_msg = "Kritische Sicherheitslücken. Sofortige Maßnahmen erforderlich."

    add_status_line(
        doc,
        "Sicherheitsstatus",
        sec_status,
        sec_color,
        extra=f"Durchschnitt: {summary['avg_security_score']}%"
    )

    # Konsistente Nachricht basierend auf dem Score
    p = doc.add_paragraph()
    run = p.add_run(sec_msg)
    run.italic = True
    run.font.color.rgb = sec_color

    doc.add_paragraph()
    doc.add_paragraph("Risiken ohne Security Headers:")
    doc.add_paragraph("• XSS-Angriffe (Cross-Site Scripting) – Einschleusung von Schadcode", style='List Bullet')
    doc.add_paragraph("• Clickjacking – Überlagern unsichtbarer Buttons auf Ihrer Seite", style='List Bullet')
    doc.add_paragraph("• Man-in-the-middle – Abfangen von Benutzerdaten", style='List Bullet')
    doc.add_paragraph("• Verringertes Benutzervertrauen (Browser zeigen Warnungen)", style='List Bullet')
    doc.add_paragraph("• Schlechteres Ranking (Google bevorzugt sichere Seiten)", style='List Bullet')
    doc.add_paragraph()

    has_security_issues = (
        summary['pages_poor_security'] > 0 or
        summary['pages_missing_security_headers'] > 0 or
        summary['pages_with_mixed_content'] > 0
    )

    # Zeige detaillierte Probleme nur wenn spezifische Issues existieren
    if has_security_issues:
        if issues['poor_security']:
            p = doc.add_paragraph()
            total = summary['pages_analyzed']
            count = len(issues['poor_security'])
            pct = round((count / total) * 100, 1) if total > 0 else 0
            run = p.add_run(f"🟠 Schwache Sicherheit: {count}/{total} Seiten ({pct}%)")
            run.bold = True
            doc.add_paragraph(
                "Seiten mit einem Security-Score <50% haben kritische Lücken bei grundlegenden Sicherheitsheadern. "
                "Unten die Endpunkte mit den größten Problemen:"
            )
            doc.add_paragraph()

            for item in issues['poor_security'][:10]:
                url = item['url']
                sec_pct = item.get('security_percentage', 0)
                missing = item.get('missing_headers', [])

                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"{url} ")
                run.font.size = Pt(9)
                run2 = p.add_run(f"({sec_pct:.1f}%)\n")
                run2.font.color.rgb = RGBColor(200, 0, 0)
                run2.font.size = Pt(9)

                if missing:
                    run3 = p.add_run(f"  Fehlt: {', '.join(missing[:4])}")
                    run3.font.size = Pt(9)
                    run3.italic = True

            if len(issues['poor_security']) > 10:
                doc.add_paragraph(f"...und {len(issues['poor_security']) - 10} weitere Seiten")
            doc.add_paragraph()

        if issues['missing_security_headers']:
            p = doc.add_paragraph()
            total = summary['pages_analyzed']
            count = len(issues['missing_security_headers'])
            pct = round((count / total) * 100, 1) if total > 0 else 0
            run = p.add_run(f"🟡 Fehlende Security Headers: {count}/{total} Seiten ({pct}%)")
            run.bold = True
            doc.add_paragraph("Seiten mit weniger als 3 Sicherheitsheadern (von 7 möglichen).")
            doc.add_paragraph()

            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("Wichtigste Security Headers (was fehlt):")
            run.bold = True
            run.font.size = Pt(11)

            security_headers_info = [
                ("HSTS (Strict-Transport-Security)",
                 "Erzwingt HTTPS-Verbindungen für eine bestimmte Zeit. Ohne: Man-in-the-middle-Angriff möglich, bei dem "
                 "Hacker über HTTP gesendete Daten abfangen. Google bevorzugt seit 2014 HTTPS-Seiten im Ranking."),
                ("CSP (Content-Security-Policy)",
                 "Legt fest, aus welchen Quellen Ressourcen (Skripte, Bilder, CSS) geladen werden dürfen. Ohne: einfache XSS-Angriffe "
                 "(Cross-Site Scripting), bei denen Hacker bösartigen JavaScript-Code in Ihre Seite einschleusen."),
                ("X-Frame-Options",
                 "Verhindert die Einbettung Ihrer Seite in iframe auf anderen Websites. Ohne: Clickjacking-Angriff – "
                 "Benutzer denkt, er klickt auf eine Schaltfläche, klickt aber tatsächlich auf eine andere."),
                ("X-Content-Type-Options",
                 "Blockiert 'MIME-Sniffing' der Browser – Erraten des Dateityps. Ohne: Browser kann eine Textdatei "
                 "als ausführbaren Code behandeln und ausführen (Angriff)."),
                ("Referrer-Policy",
                 "Kontrolliert, wie viele Informationen über die Herkunft an andere Seiten weitergegeben werden. Ohne: "
                 "die vollständige URL (mit Parametern, Tokens) kann an externe Dienste über den Referer-Header gelangen."),
                ("Permissions-Policy",
                 "Beschränkt den Zugriff auf Browser-APIs (Kamera, Mikrofon, GPS usw.). Ohne: in iframe eingebettete "
                 "Seiten können Zugriff auf sensible Benutzerressourcen anfordern."),
            ]

            for header_name, description in security_headers_info:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"{header_name}\n")
                run.bold = True
                run.font.size = Pt(10)
                run2 = p.add_run(f"  {description}")
                run2.font.size = Pt(9)
                run2.italic = True

            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("Beispiel-Endpunkte mit fehlenden Headern:")
            run.bold = True
            run.font.size = Pt(11)

            for item in issues['missing_security_headers'][:10]:
                url = item['url']
                headers_count = item['headers_count']
                missing = item.get('missing_critical', [])

                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"{url} ")
                run.font.size = Pt(9)
                run2 = p.add_run(f"({headers_count}/7 headerów)\n")
                run2.font.color.rgb = RGBColor(200, 100, 0)
                run2.font.size = Pt(9)

                if missing:
                    run3 = p.add_run(f"  Brak: {', '.join(missing[:3])}")
                    run3.font.size = Pt(9)
                    run3.italic = True

            if len(issues['missing_security_headers']) > 10:
                doc.add_paragraph(f"...oraz {len(issues['missing_security_headers']) - 10} innych stron")
            doc.add_paragraph()

        if issues['mixed_content']:
            p = doc.add_paragraph()
            total = summary['pages_analyzed']
            count = len(issues['mixed_content'])
            pct = round((count / total) * 100, 1) if total > 0 else 0
            run = p.add_run(f"🟡 Mixed content: {count}/{total} stron ({pct}%)")
            run.bold = True
            doc.add_paragraph(
                "Strony HTTPS zawierają zasoby ładowane przez HTTP (obrazki, CSS, JS). Przeglądarka pokazuje ostrzeżenie "
                "'Not Secure' mimo certyfikatu SSL. Użytkownik traci zaufanie, Google obniża ranking."
            )
            doc.add_paragraph()

            p = doc.add_paragraph()
            run = p.add_run("Strony z mixed content:")
            run.bold = True
            run.font.size = Pt(11)
            for url in issues['mixed_content'][:10]:
                doc.add_paragraph(f"• {url}", style='List Bullet')
            if len(issues['mixed_content']) > 10:
                doc.add_paragraph(f"...oraz {len(issues['mixed_content']) - 10} innych")

    if SHOW_REMEDIATIONS:
        doc.add_page_break()
        add_section_heading(doc, 'Jak wdrożyć security headers?', 2, icon='🔧')
        apache_code = '''Header always set Strict-Transport-Security "max-age=31536000; includeSubDomains"
Header always set X-Frame-Options "DENY"
Header always set X-Content-Type-Options "nosniff"
Header always set Referrer-Policy "strict-origin-when-cross-origin"
Header always set Permissions-Policy "geolocation=(), microphone=(), camera=()"'''
        q = doc.add_paragraph()
        r = q.add_run(apache_code)
        r.font.name = 'Courier New'
        r.font.size = Pt(9)
        doc.add_paragraph()
        nginx_code = '''add_header Strict-Transport-Security "max-age=31536000; includeSubDomains" always;
add_header X-Frame-Options "DENY" always;
add_header X-Content-Type-Options "nosniff" always;
add_header Referrer-Policy "strict-origin-when-cross-origin" always;
add_header Permissions-Policy "geolocation=(), microphone=(), camera=()" always;'''
        q = doc.add_paragraph()
        r = q.add_run(nginx_code)
        r.font.name = 'Courier New'
        r.font.size = Pt(9)

    doc.add_page_break()

    # =========================
    # 12. LEGENDA / SŁOWNIK
    # =========================
    add_section_heading(doc, '12. Legenda i objaśnienia – słownik dla biznesu', 1, icon='📚')

    doc.add_paragraph(
        "Poniżej znajdziesz wyjaśnienia najważniejszych terminów z raportu – językiem biznesowym, bez zbędnego żargonu IT."
    )
    doc.add_paragraph()

    # SEO
    p = doc.add_paragraph()
    run = p.add_run("SEO (Search Engine Optimization)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "To fundamenty techniczne, które decydują, czy Google w ogóle 'widzi' Twoją stronę i rozumie, o czym ona jest. "
        "SEO to m.in. tytuły stron (tag <title>), meta opisy, nagłówki H1/H2/H3, poprawne statusy HTTP, responsywność. "
        "Bez tego Google nie wie, co indeksować i gdzie Cię pokazać."
    )
    doc.add_paragraph()

    # AEO
    p = doc.add_paragraph()
    run = p.add_run("AEO (Answer Engine Optimization)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "To optymalizacja treści pod odpowiedzi na konkretne pytania użytkowników. "
        "Google i inne systemy (np. AI) starają się odpowiadać bezpośrednio w wynikach – bez klikania w link. "
        "AEO to m.in. sekcje FAQ, listy punktowane, tabelki, jasne definicje i konkrety."
    )
    doc.add_paragraph()

    # GEO
    p = doc.add_paragraph()
    run = p.add_run("GEO (Generative Engine Optimization)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "To optymalizacja pod silniki generatywne (AI), takie jak ChatGPT, Gemini, Bing Chat czy SGE Google. "
        "Te systemy nie tylko wyszukują, ale też generują odpowiedzi, opierając się na wiarygodnych źródłach. "
        "GEO to m.in. E-E-A-T, dane strukturalne Schema.org, linki do źródeł, daty publikacji i autorzy."
    )
    doc.add_paragraph()

    # E-E-A-T
    p = doc.add_paragraph()
    run = p.add_run("E-E-A-T (Experience, Expertise, Authoritativeness, Trustworthiness)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "Zestaw sygnałów jakości: doświadczenie autora, eksperckość, autorytet i zaufanie. "
        "Widoczni autorzy, daty artykułów, certyfikaty, linki do wiarygodnych źródeł, HTTPS i dane kontaktowe "
        "budują E-E-A-T i pomagają w pozycjonowaniu – szczególnie w tematach YMYL (finanse, zdrowie, prawo)."
    )
    doc.add_paragraph()

    # Schema.org
    p = doc.add_paragraph()
    run = p.add_run("Schema.org (Dane strukturalne JSON-LD)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "Ustrukturyzowany sposób opisywania zawartości strony dla Google: artykuły, produkty, FAQ, wydarzenia itp. "
        "Pozwala uzyskać rich snippets (gwiazdki, FAQ, breadcrumbs), które zwiększają CTR i widoczność."
    )
    doc.add_paragraph()

    # NAP (Local SEO)
    p = doc.add_paragraph()
    run = p.add_run("NAP (Name, Address, Phone) – Local SEO")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "Spójne dane firmy (nazwa, adres, telefon) na stronie, w Google Maps i innych serwisach. "
        "Kluczowe dla widoczności w wynikach lokalnych i Google Local Pack."
    )
    doc.add_paragraph()

    # Mobile-first indexing
    p = doc.add_paragraph()
    run = p.add_run("Mobile-first indexing")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "Google indeksuje i ocenia w pierwszej kolejności wersję mobilną strony. "
        "Jeśli wersja mobilna jest słaba, ucierpi też pozycja wersji desktopowej."
    )
    doc.add_paragraph()

    # Security Headers
    p = doc.add_paragraph()
    run = p.add_run("Security Headers (nagłówki bezpieczeństwa HTTP)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "Dodatkowe nagłówki HTTP (np. HSTS, CSP, X-Frame-Options), które chronią przed atakami i "
        "wymuszają bezpieczne zachowanie przeglądarki. Ich brak to realne ryzyko dla danych użytkowników."
    )
    doc.add_paragraph()

    # Canonical
    p = doc.add_paragraph()
    run = p.add_run("Canonical (tag kanoniczny)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "Tag <link rel='canonical'> informuje Google, która wersja URL jest główna. "
        "Chroni przed duplikacją treści i rozproszeniem 'mocy' linków między podobnymi adresami."
    )
    doc.add_paragraph()

    # Thin content
    p = doc.add_paragraph()
    run = p.add_run("Thin content (cienka treść)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "Strony z bardzo krótką, mało wartościową treścią. "
        "Google ocenia je jako low-quality, co skutkuje niższymi pozycjami i mniejszym ruchem."
    )
    doc.add_paragraph()

    # Open Graph & Twitter Cards
    p = doc.add_paragraph()
    run = p.add_run("Open Graph & Twitter Cards (meta tagi social media)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 71, 136)
    doc.add_paragraph(
        "Tagi kontrolujące wygląd linku po udostępnieniu w mediach społecznościowych (obrazek, tytuł, opis). "
        "Bez nich tracisz kontrolę nad tym, jak Twoja marka wygląda w social media."
    )
    doc.add_paragraph()

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("💡 Podsumowanie:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(50, 100, 200)
    doc.add_paragraph(
        "Każdy z opisanych elementów ma bezpośrednie przełożenie na ruch, konwersję i sprzedaż. "
        "Priorytetem są błędy krytyczne (4xx, brak Title/Description, brak viewport), następnie ostrzeżenia "
        "(canonical, schema), a na końcu optymalizacje jakościowe (E-E-A-T, content)."
    )

    doc.add_page_break()

    # =========================
    # 13. AI SUMMARY (opcjonalnie)
    # =========================
    if USE_AI_SUMMARY and OPENAI_API_KEY:
        # Zakładamy, że funkcja generate_ai_summary istnieje w innym module
        ai_summary_text = generate_ai_summary(summary, issues)  # noqa: F821
        if ai_summary_text:
            add_section_heading(doc, '13. AI-Powered Executive Summary', 1, icon='🤖')
            lines = ai_summary_text.split('\n')
            current_paragraph = None
            for line in lines:
                line = line.strip()
                if not line:
                    current_paragraph = None
                    continue
                if line.startswith('###'):
                    heading_text = line.replace('###', '').strip()
                    doc.add_heading(heading_text, level=3)
                    current_paragraph = None
                elif line.startswith('- ') or line.startswith('* '):
                    list_text = line[2:].strip()
                    doc.add_paragraph(list_text, style='List Bullet')
                    current_paragraph = None
                elif re.match(r'^\d+\.', line):
                    list_text = re.sub(r'^\d+\.\s*', '', line)
                    doc.add_paragraph(list_text, style='List Number')
                    current_paragraph = None
                else:
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph()
                    parts = re.split(r'(\*\*.*?\*\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = current_paragraph.add_run(part[2:-2])
                            run.bold = True
                        else:
                            current_paragraph.add_run(part)

    # STOPKA RAPORTU
    doc.add_paragraph('_' * 100)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Raport wygenerowany przez Audytor SEO/AEO/GEO Enhanced Edition\n')
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(120, 120, 120)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Data: {summary["generated_at"]}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)

    doc.save(word_path)
    print(f"✅ Raport Word zapisany: {word_path}")
